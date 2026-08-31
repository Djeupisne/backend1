from flask import Flask, request, jsonify, send_file, redirect
from flask_cors import CORS
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
import os, hashlib, datetime, uuid, json, re, threading, io, csv, unicodedata, zipfile, time, gc, random, tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from werkzeug.utils import secure_filename
from supabase import create_client, Client
import logging
logging.basicConfig(level=logging.INFO)
logging.getLogger('pdfminer').setLevel(logging.WARNING)
logging.getLogger('pdfplumber').setLevel(logging.WARNING)
logger = logging.getLogger(__name__)
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False
try:
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, KeepTogether
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib import fonts
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
try:
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
    logger.info("✅ OpenAI importe avec succes")
except ImportError as e:
    OPENAI_AVAILABLE = False
    logger.error(f"❌ Erreur import OpenAI: {e}")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
OPENROUTER_BASE_URL = os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1")
OPENROUTER_REASONING_ENABLED = os.getenv("OPENROUTER_REASONING_ENABLED", "false").lower() == "true"
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
logger.info(f"🔑 OPENROUTER_API_KEY: {'✅ Presente' if OPENROUTER_API_KEY else '❌ Manquante'}")
logger.info(f"🧠 OPENROUTER_REASONING: {'✅ Active' if OPENROUTER_REASONING_ENABLED else '❌ Desactive'}")
logger.info(f"🔑 DEEPSEEK_API_KEY: {'✅ Presente' if DEEPSEEK_API_KEY else '❌ Manquante'}")
_client = None
_PROVIDER = "None"
_MODEL = None
IA_ANALYSE_ACTIVE = False
ACTIVE_MODELS = []
OPENROUTER_MODELS = [
    {"name": "Nemotron 3 Ultra", "model": os.getenv("IA_MODEL_1", "nvidia/nemotron-3-ultra-550b-a55b:free"), "priority": 1, "supports_reasoning": True},
    {"name": "MiniMax M3", "model": os.getenv("IA_MODEL_2", "minimax/minimax-m3:free"), "priority": 2, "supports_reasoning": False},
    {"name": "MiniMax M2.7", "model": os.getenv("IA_MODEL_3", "minimax/minimax-m2.7:free"), "priority": 3, "supports_reasoning": False}
]
IA_FALLBACK_ENABLED = os.getenv("IA_FALLBACK_ENABLED", "true").lower() == "true"
IA_MAX_RETRIES = int(os.getenv("IA_MAX_RETRIES", "3"))
IA_MODEL_TIMEOUT = int(os.getenv("IA_MODEL_TIMEOUT", "120"))
def initialize_ia_clients():
    global _client, _PROVIDER, _MODEL, IA_ANALYSE_ACTIVE, ACTIVE_MODELS
    ACTIVE_MODELS = []
    for model_config in OPENROUTER_MODELS:
        if OPENROUTER_API_KEY:
            try:
                client = OpenAI(api_key=OPENROUTER_API_KEY, base_url=OPENROUTER_BASE_URL)
                test_response = client.chat.completions.create(
                    model=model_config["model"],
                    messages=[{"role": "user", "content": "Test"}],
                    max_tokens=5,
                    temperature=0
                )
                if test_response and test_response.choices:
                    ACTIVE_MODELS.append({
                        "client": client,
                        "model": model_config["model"],
                        "name": f"OpenRouter - {model_config['name']}",
                        "provider": "OpenRouter",
                        "base_url": OPENROUTER_BASE_URL,
                        "priority": model_config.get("priority", 10),
                        "supports_reasoning": model_config.get("supports_reasoning", False)
                    })
                    logger.info(f"✅ {model_config['name']} ({model_config['model']}) initialise avec succes")
            except Exception as e:
                logger.warning(f"⚠️ {model_config['name']} indisponible: {e}")
    if DEEPSEEK_API_KEY:
        try:
            client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
            test_response = client.chat.completions.create(
                model=DEEPSEEK_MODEL,
                messages=[{"role": "user", "content": "Test"}],
                max_tokens=5,
                temperature=0
            )
            if test_response and test_response.choices:
                ACTIVE_MODELS.append({
                    "client": client,
                    "model": DEEPSEEK_MODEL,
                    "name": "DeepSeek",
                    "provider": "DeepSeek",
                    "base_url": "https://api.deepseek.com",
                    "priority": 99,
                    "supports_reasoning": False
                })
                logger.info(f"✅ DeepSeek initialise avec succes")
        except Exception as e:
            logger.warning(f"⚠️ DeepSeek indisponible: {e}")
    ACTIVE_MODELS.sort(key=lambda x: x.get("priority", 10))
    if ACTIVE_MODELS:
        _client = ACTIVE_MODELS[0]["client"]
        _MODEL = ACTIVE_MODELS[0]["model"]
        _PROVIDER = ACTIVE_MODELS[0]["name"]
        IA_ANALYSE_ACTIVE = True
        logger.info(f"✅ {len(ACTIVE_MODELS)} modele(s) IA actif(s):")
        for m in ACTIVE_MODELS:
            supports = "🧠" if m.get("supports_reasoning", False) and OPENROUTER_REASONING_ENABLED else "📝"
            logger.info(f"   {supports} {m['name']} ({m['model']})")
    else:
        IA_ANALYSE_ACTIVE = False
        logger.warning("⚠️ AUCUN MODELE IA DISPONIBLE")
initialize_ia_clients()
_ia_semaphore = threading.Semaphore(int(os.getenv("IA_MAX_CONCURRENCY", "5")))
DOWNLOAD_MAX_RETRIES = int(os.getenv("DOWNLOAD_MAX_RETRIES", "3"))
DOWNLOAD_BASE_DELAY = float(os.getenv("DOWNLOAD_BASE_DELAY", "0.5"))
DOWNLOAD_MAX_DELAY = int(os.getenv("DOWNLOAD_MAX_DELAY", "10"))
DOWNLOAD_MAX_CONCURRENT = int(os.getenv("DOWNLOAD_MAX_CONCURRENT", "15"))
_DOWNLOAD_SEMAPHORE = threading.Semaphore(DOWNLOAD_MAX_CONCURRENT)
_ZIP_JOBS = {}
_ZIP_JOBS_LOCK = threading.Lock()
_ZIP_JOBS_MAX_AGE_SECONDS = 3600
_ZIP_MAX_WORKERS = int(os.getenv("ZIP_MAX_WORKERS", "25"))
app = Flask(__name__)
ALLOWED_ORIGINS = [
    "https://recrutment.onrender.com",
    "https://backend1-fiq5.onrender.com",
    "http://localhost:5000",
    "http://localhost:3000"
]
CORS(app, resources={
    r"/api/*": {
        "origins": ALLOWED_ORIGINS,
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization", "X-Requested-With"],
        "supports_credentials": True,
        "max_age": 600
    }
})
@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization,X-Requested-With')
    response.headers.add('Access-Control-Allow-Methods', 'GET,POST,OPTIONS,PUT,DELETE')
    response.headers.add('Access-Control-Max-Age', '600')
    if request.method == 'OPTIONS':
        response.status_code = 204
    return response
@app.route('/', methods=['GET', 'HEAD'])
def health_check():
    active_models = [{"name": m["name"], "model": m["model"], "provider": m["provider"], "supports_reasoning": m.get("supports_reasoning", False)} for m in ACTIVE_MODELS]
    return jsonify({
        'status': 'ok',
        'message': f'RecrutBank API with {len(ACTIVE_MODELS)} IA model(s)',
        'version': 'v13.1-ia-multi-fallback',
        'features': {
            'ia_available': IA_ANALYSE_ACTIVE,
            'ia_models': active_models,
            'active_models_count': len(ACTIVE_MODELS),
            'analysis_method': '100% IA avec fallback multi-modeles',
            'scoring_max': 14,
            'postes_actifs': ["Data Analyst Finance"],
            'postes_clotures': ["Chef de Division Local Corporate"],
            'version': 'v13.1-ia-multi-fallback'
        }
    }), 200
app.config['JWT_SECRET_KEY'] = os.getenv("JWT_SECRET_KEY", "gestion-candidatures-secret-2024")
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = datetime.timedelta(hours=8)
jwt = JWTManager(app)
SUPABASE_URL = os.getenv("SUPABASE_URL", "")
SUPABASE_KEY = os.getenv("SUPABASE_KEY", "")
SUPABASE_STORAGE_BUCKET = os.getenv("SUPABASE_STORAGE_BUCKET", "candidatures")
supabase = create_client(SUPABASE_URL, SUPABASE_KEY) if SUPABASE_URL and SUPABASE_KEY else None
app.config['SMTP_HOST'] = os.getenv('SMTP_HOST', 'smtp.gmail.com')
app.config['SMTP_PORT'] = int(os.getenv('SMTP_PORT', 587))
app.config['SMTP_USER'] = os.getenv('SMTP_USER', '')
app.config['SMTP_PASSWORD'] = os.getenv('SMTP_PASSWORD', '')
app.config['SMTP_FROM'] = os.getenv('SMTP_FROM', 'RecrutBank RH <oualoumidjeupisne@gmail.com>')
app.config['SMTP_USE_TLS'] = os.getenv('SMTP_USE_TLS', 'true').lower() == 'true'
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt'}
app.config['MAX_CONTENT_LENGTH'] = 15 * 1024 * 1024
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
def upload_file_to_supabase(file_obj, blob_name, content_type=None):
    if not supabase:
        return None
    try:
        file_bytes = file_obj.read()
        supabase.storage.from_(SUPABASE_STORAGE_BUCKET).upload(
            blob_name,
            file_bytes,
            {"content-type": content_type or "application/octet-stream", "upsert": "true"}
        )
        return blob_name
    except Exception as e:
        logger.error(f"Upload error: {e}")
        return None
def retry_with_backoff(max_retries=DOWNLOAD_MAX_RETRIES, base_delay=DOWNLOAD_BASE_DELAY, max_delay=DOWNLOAD_MAX_DELAY):
    def decorator(func):
        def wrapper(*args, **kwargs):
            last_exception = None
            for attempt in range(max_retries):
                try:
                    result = func(*args, **kwargs)
                    return result
                except Exception as e:
                    last_exception = e
                    error_str = str(e).lower()
                    retryable_keywords = ["errno 11", "resource temporarily unavailable", "timeout", "connection", "temporarily unavailable", "rate limit", "too many requests", "503", "502", "504", "connection refused", "connection reset"]
                    if not any(kw in error_str for kw in retryable_keywords):
                        raise
                    if attempt == max_retries - 1:
                        raise
                    delay = min(base_delay * (2 ** attempt), max_delay)
                    jitter = random.uniform(0, delay * 0.3)
                    total_delay = delay + jitter
                    logger.warning(f"Tentative {attempt + 1}/{max_retries} echouee: {e}. Nouvel essai dans {total_delay:.2f}s")
                    time.sleep(total_delay)
            raise last_exception
        return wrapper
    return decorator
@retry_with_backoff(max_retries=DOWNLOAD_MAX_RETRIES)
def download_file_from_supabase_robust(blob_name):
    if not supabase:
        return None
    with _DOWNLOAD_SEMAPHORE:
        try:
            response = supabase.storage.from_(SUPABASE_STORAGE_BUCKET).download(blob_name)
            return response
        except Exception as e:
            logger.error(f"Erreur telechargement {blob_name}: {e}")
            raise
def download_file_from_supabase(blob_name):
    if not supabase:
        return None
    try:
        response = supabase.storage.from_(SUPABASE_STORAGE_BUCKET).download(blob_name)
        return response
    except Exception as e:
        error_str = str(e).lower()
        if any(kw in error_str for kw in ["errno 11", "resource temporarily unavailable", "timeout", "connection"]):
            return download_file_from_supabase_robust(blob_name)
        logger.error(f"Download error: {e}")
        return None
def get_signed_url(blob_name, expiration_minutes=60):
    if not supabase:
        return None
    try:
        response = supabase.storage.from_(SUPABASE_STORAGE_BUCKET).create_signed_url(blob_name, expiration_minutes * 60)
        return response.get('signedURL') if response else None
    except Exception as e:
        logger.error(f"Signed URL error: {e}")
        return None
def send_email(to_email, subject, body):
    import requests
    import re as _re
    brevo_api_key = os.getenv('BREVO_API_KEY', '')
    smtp_from = os.getenv('SMTP_FROM', 'RecrutBank RH <oualoumidjeupisne@gmail.com>')
    if not brevo_api_key:
        return False
    match = _re.search(r'<(.+?)>', smtp_from)
    sender_email = match.group(1) if match else smtp_from
    sender_name = smtp_from.split('<')[0].strip() if '<' in smtp_from else 'RecrutBank RH'
    html_content = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"></head><body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;"><div style="max-width: 600px; margin: 0 auto; padding: 20px;">{body.replace(chr(10), '<br>')}</div></body></html>"""
    url = "https://api.brevo.com/v3/smtp/email"
    headers = {"api-key": brevo_api_key, "Content-Type": "application/json", "Accept": "application/json"}
    payload = {
        "sender": {"name": sender_name, "email": sender_email},
        "to": [{"email": to_email, "name": to_email.split('@')[0]}],
        "subject": subject,
        "htmlContent": html_content,
        "textContent": body
    }
    try:
        response = requests.post(url, json=payload, headers=headers, timeout=30)
        return response.status_code == 201
    except Exception:
        return False
def hash_pwd(pwd):
    return hashlib.sha256(pwd.encode()).hexdigest()
def normalize_spaces(text):
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text)
    return text.strip()
def normalize_unicode(text):
    if not text:
        return ""
    text = unicodedata.normalize('NFC', text)
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    text = re.sub(r'[\u00A0\u1680\u2000-\u200B\u2028\u2029\u202F\u205F\u3000]', ' ', text)
    return text.strip()
MAX_PDF_PAGES = 10
MAX_PDF_SIZE_BYTES = 10 * 1024 * 1024
MAX_TEXT_SIZE = 15000
def extract_text_from_pdf_robust(file_bytes, filename):
    if len(file_bytes) > MAX_PDF_SIZE_BYTES:
        logger.warning(f"PDF trop volumineux: {filename}")
        return ""
    text = ""
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                total_pages = min(len(pdf.pages), MAX_PDF_PAGES)
                for i in range(total_pages):
                    try:
                        page = pdf.pages[i]
                        content = page.extract_text(x_tolerance=3, y_tolerance=3, keep_blank_chars=True, use_text_flow=True)
                        if content:
                            text += normalize_spaces(content) + "\n"
                        if len(text) > MAX_TEXT_SIZE:
                            text = text[:MAX_TEXT_SIZE]
                            break
                    except Exception as e:
                        logger.warning(f"pdfplumber page {i} erreur: {e}")
                        continue
            if text.strip() and len(text.strip()) > 50:
                return normalize_unicode(text.strip())
        except Exception as e:
            logger.warning(f"pdfplumber erreur: {e}")
    if PYPDF2_AVAILABLE:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            total_pages = min(len(reader.pages), MAX_PDF_PAGES)
            for i in range(total_pages):
                try:
                    content = reader.pages[i].extract_text()
                    if content:
                        text += normalize_spaces(content) + "\n"
                    if len(text) > MAX_TEXT_SIZE:
                        text = text[:MAX_TEXT_SIZE]
                        break
                except Exception as e:
                    logger.warning(f"PyPDF2 page {i} erreur: {e}")
                    continue
            if text.strip() and len(text.strip()) > 50:
                return normalize_unicode(text.strip())
        except Exception as e:
            logger.warning(f"PyPDF2 erreur: {e}")
    return text.strip() if text.strip() else ""
def extract_text_from_docx_robust(file_bytes):
    if not DOCX_AVAILABLE:
        return ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            t = normalize_spaces(para.text)
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    ct = normalize_spaces(cell.text)
                    if ct:
                        cells.append(ct)
                if cells:
                    parts.append(" | ".join(cells))
        result = "\n".join(parts).strip()
        if len(result) > MAX_TEXT_SIZE:
            result = result[:MAX_TEXT_SIZE]
        if len(result) > 50:
            return normalize_unicode(result)
    except Exception as e:
        logger.warning(f"DOCX extraction avec python-docx echouee: {e}")
    try:
        import struct
        text = ""
        for i in range(0, len(file_bytes), 2):
            if i + 1 < len(file_bytes):
                try:
                    char = struct.unpack('<H', file_bytes[i:i+2])[0]
                    if 32 <= char <= 126 or char in [10, 13]:
                        text += chr(char)
                    else:
                        text += ' '
                except:
                    text += ' '
        text = re.sub(r'\s+', ' ', text).strip()
        if len(text) > 50:
            return normalize_unicode(text[:MAX_TEXT_SIZE])
    except:
        pass
    try:
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
            try:
                text = file_bytes.decode(encoding, errors='ignore')
                text = re.sub(r'[^\w\s\.\,\:\;\-\?\!\@\#\%\&\*\(\)\-\+\=\/\'\"]', ' ', text)
                text = re.sub(r'\s+', ' ', text).strip()
                if len(text) > 50:
                    return normalize_unicode(text[:MAX_TEXT_SIZE])
            except:
                continue
    except Exception as e:
        logger.warning(f"Extraction brute echouee: {e}")
    return ""
def extract_text_from_txt(file_bytes):
    if CHARDET_AVAILABLE:
        try:
            detected = chardet.detect(file_bytes[:10000])
            encoding = detected['encoding'] or 'utf-8'
            text = file_bytes.decode(encoding, errors='ignore')
            if len(text) > MAX_TEXT_SIZE:
                text = text[:MAX_TEXT_SIZE]
            return normalize_unicode(normalize_spaces(text))
        except Exception:
            pass
    for enc in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']:
        try:
            text = file_bytes.decode(enc, errors='ignore').strip()
            if len(text) > MAX_TEXT_SIZE:
                text = text[:MAX_TEXT_SIZE]
            return normalize_unicode(normalize_spaces(text))
        except (UnicodeDecodeError, UnicodeError):
            continue
    return ""
def extract_text_robust_from_bytes(file_bytes, filename):
    if not file_bytes:
        return ""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    text = ""
    if ext == 'pdf':
        text = extract_text_from_pdf_robust(file_bytes, filename)
        logger.info(f"Extraction PDF {filename}: {len(text)} caracteres")
    elif ext in ('doc', 'docx'):
        text = extract_text_from_docx_robust(file_bytes)
        logger.info(f"Extraction DOC/DOCX {filename}: {len(text)} caracteres")
    elif ext == 'txt':
        text = extract_text_from_txt(file_bytes)
        logger.info(f"Extraction TXT {filename}: {len(text)} caracteres")
    return text.strip() if text.strip() else ""
def init_recruteur():
    try:
        if supabase:
            response = supabase.table('recruteurs').select('*').eq('email', 'sougnabeoualoumibank@gmail.com').execute()
            if not response.data:
                supabase.table('recruteurs').insert({
                    "email": "sougnabeoualoumibank@gmail.com",
                    "password": hash_pwd("AdminLaurent123"),
                    "nom": "Responsable RH"
                }).execute()
    except Exception as e:
        logger.warning(f"Erreur initialisation recruteur : {e}")
init_recruteur()
POSTES = [
    "Chef de Division Local Corporate",
    "Data Analyst Finance",
    "Responsable Administration de Credit",
    "Analyste Credit CCB",
    "Archiviste (Administration Credit)",
    "Senior Finance Officer",
    "Market Risk Officer",
    "IT Reseau & Infrastructure",
    "Auditeur interne",
    "Chef service controle des engagements",
    "Chef service IT (maintenance/support)",
    "Chef service finance",
    "Chef service risques de marche",
    "Chef service reporting reglementaire",
    "Chef de Section Compensation",
    "Charge(e) d'Administration de Credit"
]
POSTES_ACTIFS = ["Data Analyst Finance"]
POSTES_CLOTURES = [p for p in POSTES if p not in POSTES_ACTIFS]
def is_poste_actif(poste):
    return poste in POSTES_ACTIFS
GRILLE = {
    "Chef de Division Local Corporate": {
        "eliminatoire": [
            "Aucune experience dans le secteur bancaire ou financier reglemente",
            "Niveau de diplome inferieur a Bac +4 (Master ou equivalent requis)",
            "Moins de 5 ans d'experience professionnelle, dont une partie significative en banque",
            "Aucune experience en gestion d'un portefeuille de clients SME (PME)/local corporate ou d'entreprises",
            "Aucune experience manageriale : ni encadrement d'equipe, ni pilotage d'une activite commerciale",
            "Aucune exposition a la gestion du risque de credit ou au suivi de la qualite d'un portefeuille (NPL, provisions)"
        ],
        "a_verifier": [
            "A pilote une activite Local Corporate/SME (PME) avec des objectifs chiffres (revenus, volumes, marges)",
            "A gere un portefeuille de clients Local Corporate/SME (PME) et demontre sa capacite a le developper",
            "A encadre et evalue une equipe commerciale ou bancaire",
            "A assure le suivi de la qualite du portefeuille de credit (NPL, CIR, provisions) et rendu compte a la direction",
            "A developpe des ventes croisees (cross-selling) ou des partenariats interdépartementaux",
            "A produit ou supervise des rapports de performance commerciale et financiere",
            "A une exposition a la reglementation bancaire locale (COBAC, BEAC) ou internationale"
        ],
        "signaux_forts": [
            "Pilotage d'une division ou d'une ligne Local Corporate/SME (PME) avec atteinte des objectifs de revenus et de portefeuille",
            "Gestion active du ratio NPL (creances douteuses) et du ratio cout/revenu (CIR) - resultats chiffres mentionnes",
            "Experience averee en cross-selling avec des equipes TSG, Trade Finance ou Cash Management",
            "Developpement reel du portefeuille Local Corporate : acquisition de clients, fidelisation, nombre de produits par client",
            "Leadership demontre : constitution d'equipe, developpement des collaborateurs, vivier de talents",
            "Certification Ecobank / Moody's ou ITB (Institut Technique de Banque) ou equivalent",
            "Connaissance du marche Local Corporate/SME (PME) tchadien ou de la zone CEMAC/UEMOA",
            "Exposition aux plateformes numeriques bancaires (OMNI, Cash Management ou equivalent)",
            "Resultats commerciaux quantifies et verifiables dans son CV (chiffres d'affaires, taux de croissance, NPS)"
        ],
        "points_attention": [
            "Parcours exclusivement back-office ou risques sans experience commerciale Local Corporate/SME (PME)",
            "Profil techniquement solide (credit, analyse) mais sans experience manageriale ni pilotage de P&L",
            "Experiences tres courtes (moins de 2 ans par poste) sans progression hierarchique visible",
            "CV sans resultats chiffres (missions decrites en responsabilites sans livrables ni indicateurs atteints)",
            "Mobilite geographique ou sectorielle excessive sans ancrage dans le secteur bancaire Local Corporate",
            "Trous inexpliques dans le parcours ou incoherences entre les postes declares"
        ],
        "scores_max": {
            "experience_corporate": 3,
            "management": 3,
            "risque_credit": 2,
            "cross_selling": 2,
            "coherence_parcours": 2,
            "qualite_cv": 1,
            "certification": 1
        },
        "regles_metier": [
            "Chef d'agence / Directeur d'agence : valide automatiquement le management, le portefeuille commercial et le risque de credit (NPL/provisions).",
            "Gestionnaire de portefeuille / Relationship Manager SME/Corporate : valide automatiquement l'exposition au risque de credit et la gestion des impayes.",
            "Bac+3 compensable par 10+ annees d'experience averee en SME/Local Corporate."
        ],
        "score_max_total": 14,
        "description": "Chef de Division Local Corporate - Pilotage d'une activite commerciale SME/Corporate"
    },
    "Data Analyst Finance": {
        "eliminatoire": [
            "Aucune formation en Finance, Comptabilite, Controle de gestion, Statistiques, Data Analytics ou Informatique decisionnelle",
            "Niveau de diplome inferieur a Bac +3",
            "Aucune experience en analyse financiere, reporting financier, controle de gestion, audit ou data analytics",
            "Aucune maitrise d'Excel (TCD, formules, Power Query) - competence incontournable et non negociable",
            "Aucune connaissance de la comptabilite ou des etats financiers (P&L, bilan, flux de tresorerie)"
        ],
        "a_verifier": [
            "A produit des rapports financiers periodiques (mensuels, trimestriels)",
            "A conçu ou maintenu des tableaux de bord financiers (Power BI, Excel ou autre outil BI)",
            "A realise des analyses Budget / Realise / N-1 avec identification des ecarts",
            "A travaille avec SQL pour extraire ou interroger des donnees financieres",
            "A assure la reconciliation de donnees multi-sources (comptabilite / systemes operationnels)",
            "A participe a l'elaboration d'un budget ou d'un forecast financier",
            "A une experience dans le secteur bancaire ou avec un Core Banking (FLEXCUBE, T24, Amplitude)"
        ],
        "signaux_forts": [
            "Maitrise explicite de Power BI (dashboards, DAX, Power Query) avec exemples de realisations concretes",
            "Experience averee en automatisation de reportings (Power Query, VBA, Python, outils ETL)",
            "Analyse d'ecarts Budget / Realise / N-1 avec presentation a la Direction Financiere ou a la DG",
            "Participation a la construction de modeles de prevision financiere ou d'analyses de scenarios",
            "Exposition aux donnees bancaires : PNB, NPL, cout du risque, rentabilite par agence ou produit",
            "Maitrise de SQL pour l'extraction et la manipulation de donnees en base relationnelle",
            "Connaissance de Python ou R pour des analyses statistiques avancees",
            "Mise en place de controles qualite sur les donnees et documentation des regles de calcul",
            "Resultats quantifies dans le CV : gains de productivite, delais reduits, anomalies detectees"
        ],
        "points_attention": [
            "Profil purement comptable sans exposition aux outils BI ou au reporting de gestion",
            "Profil exclusivement IT / developpeur sans connaissance financiere",
            "Experience uniquement academique ou stage sans production de reportings reels en environnement professionnel",
            "CV sans aucun outil cite nommement",
            "Missions decrites en termes generiques sans livrables precis ni resultats mesurables",
            "Trous inexpliques dans le parcours ou experiences tres courtes sans progression visible"
        ],
        "scores_max": {
            "experience_finance_data": 3,
            "outils_bi": 3,
            "sql": 2,
            "exposition_bancaire": 2,
            "coherence_parcours": 2,
            "qualite_cv": 1,
            "competences_avancees": 1
        },
        "regles_metier": [
            "Bac+3 compensable par 4-5 ans d'experience averee en analyse financiere et BI.",
            "La maitrise d'Excel (TCD, Power Query) est un pre-requis absolu et non negociable."
        ],
        "score_max_total": 14,
        "description": "Data Analyst Finance - Analyse de donnees financieres et reporting"
    }
}
GRILLE["Responsable Administration de Credit"] = {
    "eliminatoire": [
        "Aucune experience en administration de credit ou dans un service credit",
        "Niveau de diplome inferieur a Bac +3",
        "Moins de 3 ans d'experience dans le secteur bancaire",
        "Aucune connaissance des processus de montage et d'instruction de credit"
    ],
    "a_verifier": [
        "A instruit des dossiers de credit en respectant les procedures",
        "A assure la mise en place et le suivi des garanties",
        "A travaille avec un Core Banking (FLEXCUBE, T24)",
        "A participe a l'elaboration des rapports de suivi de portefeuille",
        "A une connaissance de la reglementation bancaire (COBAC)"
    ],
    "signaux_forts": [
        "Experience dans un service Administration de Credit",
        "Maitrise des outils de gestion de credit",
        "Connaissance des processus de recouvrement",
        "Certification en gestion de credit ou analyse financiere"
    ],
    "points_attention": [
        "Profil sans experience bancaire",
        "Experience uniquement en recouvrement sans instruction de credit",
        "CV sans precision sur le type de credits traites"
    ],
    "scores_max": {
        "experience_credit": 3,
        "management": 1,
        "risque_credit": 3,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 11,
    "description": "Responsable Administration de Credit - Gestion des dossiers de credit"
}
GRILLE["Analyste Credit CCB"] = {
    "eliminatoire": [
        "Aucune experience en analyse de credit ou dans une banque",
        "Niveau de diplome inferieur a Bac +3",
        "Aucune connaissance de l'analyse financiere d'entreprises"
    ],
    "a_verifier": [
        "A realise des analyses financieres de dossiers de credit",
        "A produit des rapports d'analyse et des recommandations",
        "A travaille avec des outils d'analyse de credit",
        "A une connaissance des normes IFRS et de la reglementation"
    ],
    "signaux_forts": [
        "Experience en analyse de risque de credit",
        "Certification en analyse financiere",
        "Maitrise des ratios financiers et de l'analyse de bilan",
        "Connaissance du secteur SME/PME"
    ],
    "points_attention": [
        "Profil sans experience en analyse credit",
        "Experience uniquement en back-office sans analyse",
        "CV sans mention des methodes d'analyse utilisees"
    ],
    "scores_max": {
        "experience_credit": 3,
        "management": 0,
        "risque_credit": 3,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 10,
    "description": "Analyste Credit CCB - Analyse de risque de credit"
}
GRILLE["Archiviste (Administration Credit)"] = {
    "eliminatoire": [
        "Aucune experience en archivage ou gestion documentaire",
        "Niveau de diplome inferieur a Bac +2"
    ],
    "a_verifier": [
        "A experience en gestion de dossiers physiques et electroniques",
        "Maitrise des outils de GED (Gestion Electronique de Documents)",
        "Connaissance des procedures de classification et d'archivage",
        "Experience dans le secteur bancaire ou financier"
    ],
    "signaux_forts": [
        "Certification en archivage ou GED",
        "Experience dans un service bancaire",
        "Connaissance des normes d'archivage (ISO)",
        "Maitrise de l'outil informatique"
    ],
    "points_attention": [
        "Experience uniquement en saisie sans gestion d'archives",
        "CV sans mention des outils utilises",
        "Parcours sans progression"
    ],
    "scores_max": {
        "experience_archivage": 3,
        "management": 0,
        "risque_credit": 0,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 7,
    "description": "Archiviste - Gestion documentaire"
}
GRILLE["Senior Finance Officer"] = {
    "eliminatoire": [
        "Aucune experience en finance d'entreprise ou reporting financier",
        "Niveau de diplome inferieur a Bac +4",
        "Moins de 5 ans d'experience en finance"
    ],
    "a_verifier": [
        "A produit des rapports financiers periodiques",
        "A participe a l'elaboration des budgets",
        "A assure le suivi de la tresorerie",
        "A travaille avec des outils de gestion financiere (SAP, Oracle)",
        "A une experience en controle de gestion"
    ],
    "signaux_forts": [
        "Certification en finance (CFA, CPA)",
        "Experience en banque ou institution financiere",
        "Maitrise des normes IFRS",
        "Experience en audit ou controle interne"
    ],
    "points_attention": [
        "Profil sans experience bancaire",
        "Experience uniquement en comptabilite sans reporting",
        "CV sans mention des outils et methodes"
    ],
    "scores_max": {
        "experience_finance": 3,
        "management": 2,
        "risque_credit": 0,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 9,
    "description": "Senior Finance Officer - Gestion financiere et reporting"
}
GRILLE["Market Risk Officer"] = {
    "eliminatoire": [
        "Aucune experience en gestion des risques de marche",
        "Niveau de diplome inferieur a Bac +4",
        "Aucune connaissance des produits financiers et des risques associes"
    ],
    "a_verifier": [
        "A evalue les risques de marche (taux, change, actions)",
        "A produit des rapports de risque (VaR, stress tests)",
        "A travaille avec des outils de risque (Murex, RiskMetrics)",
        "A une connaissance des normes Bâle III"
    ],
    "signaux_forts": [
        "Certification PRM (Professional Risk Manager)",
        "Experience en banque d'investissement",
        "Maitrise des modeles de risque avances",
        "Connaissance de la reglementation COBAC"
    ],
    "points_attention": [
        "Profil sans experience en risque de marche",
        "Experience uniquement en risque de credit",
        "CV sans mention des outils et methodes"
    ],
    "scores_max": {
        "experience_risque": 3,
        "management": 1,
        "risque_credit": 0,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 8,
    "description": "Market Risk Officer - Gestion des risques de marche"
}
GRILLE["IT Reseau & Infrastructure"] = {
    "eliminatoire": [
        "Aucune experience en administration reseau ou infrastructure IT",
        "Niveau de diplome inferieur a Bac +3",
        "Aucune certification IT reconnue"
    ],
    "a_verifier": [
        "A administre des reseaux (Cisco, Juniper)",
        "A gere des infrastructures serveurs (Windows, Linux)",
        "A assure la securite des reseaux",
        "A experience en virtualisation (VMware, Hyper-V)",
        "A travaille dans le secteur bancaire ou financier"
    ],
    "signaux_forts": [
        "Certifications Cisco (CCNA, CCNP)",
        "Certifications Microsoft",
        "Experience en Cybersecurite",
        "Connaissance des normes de securite bancaire"
    ],
    "points_attention": [
        "Profil sans certification",
        "Experience uniquement en support sans administration",
        "CV sans mention des technologies maitrisees"
    ],
    "scores_max": {
        "experience_it": 3,
        "management": 2,
        "risque_credit": 0,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 9,
    "description": "IT Reseau & Infrastructure - Gestion des infrastructures IT"
}
GRILLE["Auditeur interne"] = {
    "eliminatoire": [
        "Aucune experience en audit interne ou externe",
        "Niveau de diplome inferieur a Bac +4",
        "Aucune certification en audit"
    ],
    "a_verifier": [
        "A realise des missions d'audit",
        "A produit des rapports d'audit avec recommandations",
        "A une connaissance des normes d'audit",
        "A travaille dans le secteur bancaire",
        "A une connaissance de la reglementation"
    ],
    "signaux_forts": [
        "Certification CIA (Certified Internal Auditor)",
        "Experience en banque",
        "Connaissance des normes internationales d'audit",
        "Experience en audit des risques operationnels"
    ],
    "points_attention": [
        "Profil sans experience en audit",
        "Experience uniquement en comptabilite sans audit",
        "CV sans mention des missions d'audit"
    ],
    "scores_max": {
        "experience_audit": 3,
        "management": 2,
        "risque_credit": 0,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 9,
    "description": "Auditeur interne - Audit et controle des risques"
}
GRILLE["Chef service controle des engagements"] = {
    "eliminatoire": [
        "Aucune experience en controle des engagements ou risque de credit",
        "Niveau de diplome inferieur a Bac +4",
        "Moins de 5 ans d'experience dans le secteur bancaire"
    ],
    "a_verifier": [
        "A supervise des activites de controle des engagements",
        "A assure le suivi du portefeuille de credits",
        "A produit des rapports de controle et de conformite",
        "A une connaissance de la reglementation bancaire",
        "A encadre une equipe"
    ],
    "signaux_forts": [
        "Experience en management d'equipe",
        "Connaissance des normes COBAC",
        "Certification en risque de credit",
        "Experience dans une banque internationale"
    ],
    "points_attention": [
        "Profil sans experience en controle",
        "Experience uniquement en analyse sans controle",
        "CV sans mention des indicateurs de suivi"
    ],
    "scores_max": {
        "experience_credit": 3,
        "management": 3,
        "risque_credit": 3,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 13,
    "description": "Chef service controle des engagements - Supervision du controle credit"
}
GRILLE["Chef service IT (maintenance/support)"] = {
    "eliminatoire": [
        "Aucune experience en maintenance IT ou support",
        "Niveau de diplome inferieur a Bac +3",
        "Aucune experience en management d'equipe"
    ],
    "a_verifier": [
        "A supervise une equipe de maintenance IT",
        "A assure la gestion des incidents et des demandes",
        "A mis en place des procedures de maintenance",
        "A travaille avec des outils de ticketing",
        "A une connaissance des normes ITIL"
    ],
    "signaux_forts": [
        "Certification ITIL",
        "Experience en banque",
        "Connaissance des systemes bancaires (Core Banking)",
        "Management d'equipe IT"
    ],
    "points_attention": [
        "Profil sans experience en management",
        "Experience uniquement en developpement sans maintenance",
        "CV sans mention des outils de gestion"
    ],
    "scores_max": {
        "experience_it": 3,
        "management": 3,
        "risque_credit": 0,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 10,
    "description": "Chef service IT - Supervision des activites IT"
}
GRILLE["Chef service finance"] = {
    "eliminatoire": [
        "Aucune experience en finance ou comptabilite",
        "Niveau de diplome inferieur a Bac +4",
        "Moins de 5 ans d'experience en finance",
        "Aucune experience en management"
    ],
    "a_verifier": [
        "A supervise les activites financieres",
        "A produit des rapports financiers",
        "A assure le reporting a la direction",
        "A participe a l'elaboration des budgets",
        "A une connaissance de la reglementation bancaire"
    ],
    "signaux_forts": [
        "Certification en finance (CFA, CPA)",
        "Experience en banque",
        "Maitrise des normes IFRS",
        "Connaissance des Core Banking",
        "Management d'equipe financiere"
    ],
    "points_attention": [
        "Profil sans experience en management",
        "Experience uniquement en comptabilite sans reporting",
        "CV sans mention des outils et methodes"
    ],
    "scores_max": {
        "experience_finance": 3,
        "management": 3,
        "risque_credit": 0,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 10,
    "description": "Chef service finance - Supervision des activites financieres"
}
GRILLE["Chef service risques de marche"] = {
    "eliminatoire": [
        "Aucune experience en gestion des risques de marche",
        "Niveau de diplome inferieur a Bac +4",
        "Moins de 5 ans d'experience dans le secteur bancaire",
        "Aucune experience en management"
    ],
    "a_verifier": [
        "A supervise les activites de gestion des risques de marche",
        "A produit des rapports de risque (VaR, stress tests)",
        "A assure le reporting reglementaire",
        "A une connaissance des normes Bâle III",
        "A encadre une equipe"
    ],
    "signaux_forts": [
        "Certification PRM ou FRM",
        "Experience en banque d'investissement",
        "Connaissance de la reglementation COBAC",
        "Management d'equipe risque",
        "Experience en modelisation financiere"
    ],
    "points_attention": [
        "Profil sans experience en risque de marche",
        "Experience uniquement en risque de credit",
        "CV sans mention des methodes de calcul"
    ],
    "scores_max": {
        "experience_risque": 3,
        "management": 3,
        "risque_credit": 0,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 10,
    "description": "Chef service risques de marche - Supervision des risques de marche"
}
GRILLE["Chef service reporting reglementaire"] = {
    "eliminatoire": [
        "Aucune experience en reporting reglementaire bancaire",
        "Niveau de diplome inferieur a Bac +4",
        "Moins de 5 ans d'experience dans le secteur bancaire",
        "Aucune experience en management"
    ],
    "a_verifier": [
        "A supervise la production de reportings reglementaires",
        "A assure la conformite aux normes COBAC/BEAC",
        "A produit des rapports pour les autorites de tutelle",
        "A une connaissance des normes prudentielles",
        "A encadre une equipe"
    ],
    "signaux_forts": [
        "Connaissance approfondie de la reglementation COBAC",
        "Experience en banque",
        "Certification en conformite bancaire",
        "Management d'equipe reporting",
        "Maitrise des outils de reporting"
    ],
    "points_attention": [
        "Profil sans experience en reporting reglementaire",
        "Experience uniquement en reporting interne",
        "CV sans mention des normes et reglementations"
    ],
    "scores_max": {
        "experience_reporting": 3,
        "management": 3,
        "risque_credit": 0,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 10,
    "description": "Chef service reporting reglementaire - Supervision du reporting reglementaire"
}
GRILLE["Chef de Section Compensation"] = {
    "eliminatoire": [
        "Aucune experience en compensation ou operations bancaires",
        "Niveau de diplome inferieur a Bac +3",
        "Aucune experience en management"
    ],
    "a_verifier": [
        "A supervise les activites de compensation",
        "A assure la reconciliation des operations",
        "A travaille avec un Core Banking",
        "A une connaissance des procedures bancaires",
        "A encadre une equipe"
    ],
    "signaux_forts": [
        "Experience en banque",
        "Connaissance des systemes de compensation",
        "Management d'equipe operations",
        "Maitrise des outils bancaires"
    ],
    "points_attention": [
        "Profil sans experience en management",
        "Experience uniquement en back-office sans supervision",
        "CV sans mention des systemes utilises"
    ],
    "scores_max": {
        "experience_operations": 3,
        "management": 3,
        "risque_credit": 0,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 10,
    "description": "Chef de Section Compensation - Supervision des activites de compensation"
}
GRILLE["Charge(e) d'Administration de Credit"] = {
    "eliminatoire": [
        "Aucune experience en administration de credit",
        "Niveau de diplome inferieur a Bac +3"
    ],
    "a_verifier": [
        "A instruit des dossiers de credit",
        "A assure le suivi administratif des dossiers",
        "A travaille avec un Core Banking",
        "A une connaissance des procedures de credit",
        "A assure le suivi des garanties"
    ],
    "signaux_forts": [
        "Experience dans un service credit bancaire",
        "Connaissance des processus de credit",
        "Maitrise des outils de gestion de credit",
        "Autonomie et rigueur demontrees"
    ],
    "points_attention": [
        "Profil sans experience bancaire",
        "Experience uniquement en recouvrement",
        "CV sans precision sur les types de credits traites"
    ],
    "scores_max": {
        "experience_credit": 3,
        "management": 0,
        "risque_credit": 2,
        "cross_selling": 0,
        "coherence_parcours": 2,
        "qualite_cv": 1,
        "certification": 1
    },
    "regles_metier": [],
    "score_max_total": 9,
    "description": "Charge(e) d'Administration de Credit - Gestion des dossiers de credit"
}
def parse_json_robust(result_text):
    import re as re_json
    if not result_text:
        return None
    try:
        analyse = json.loads(result_text.strip())
        logger.info("✅ JSON parse avec succes")
        return analyse
    except json.JSONDecodeError as e:
        logger.warning(f"⚠️ Erreur parsing JSON direct: {e}")
        cleaned_text = result_text
        cleaned_text = re_json.sub(r'//.*?$', '', cleaned_text, flags=re_json.MULTILINE)
        cleaned_text = re_json.sub(r',(\s*[}\]])', r'\1', cleaned_text)
        cleaned_text = re_json.sub(r'[\x00-\x1F\x7F]', '', cleaned_text)
        cleaned_text = re_json.sub(r'```json\s*', '', cleaned_text)
        cleaned_text = re_json.sub(r'\s*```', '', cleaned_text)
        cleaned_text = re_json.sub(r'^[^{]*', '', cleaned_text)
        cleaned_text = re_json.sub(r'[^}]*$', '', cleaned_text)
        try:
            analyse = json.loads(cleaned_text)
            logger.info("✅ JSON parse apres nettoyage")
            return analyse
        except json.JSONDecodeError as e2:
            logger.warning(f"⚠️ Erreur parsing JSON apres nettoyage: {e2}")
            json_patterns = [
                r'(?:```json\s*)?(\{[\s\S]*?\})(?:\s*```)?',
                r'(\{[\s\S]*\})',
                r'(\[[\s\S]*\])'
            ]
            for pattern in json_patterns:
                matches = re_json.findall(pattern, cleaned_text)
                if matches:
                    for match in matches:
                        try:
                            parsed = json.loads(match.strip())
                            if isinstance(parsed, (dict, list)):
                                logger.info("✅ JSON extrait avec regex")
                                return parsed
                        except json.JSONDecodeError:
                            continue
            logger.error("❌ Aucun JSON trouve dans la reponse")
            return None
def get_recommandation_from_score(score, flags_elim=None, score_max=14):
    try:
        s = float(score) if score is not None else 0
    except (ValueError, TypeError):
        s = 0
    if flags_elim and len(flags_elim) > 0:
        return "Rejet - Critere(s) eliminatoire(s) non satisfait(s)"
    if s >= 11:
        return "Entretien prioritaire"
    elif s >= 7:
        return "Potentiel a evaluer en entretien"
    else:
        return "Rejet"
def get_statut_from_decision(decision, flags_elim=None):
    if flags_elim and len(flags_elim) > 0:
        return "rejete"
    if not decision:
        return "en_attente"
    if "Entretien prioritaire" in decision or "Shortlist" in decision:
        return "retenu"
    elif "Potentiel" in decision or "considerer" in decision or "Faible" in decision:
        return "entretien"
    else:
        return "rejete"
def sort_candidats(candidats):
    """
    Trie les candidats par :
    1. Statut (retenu en premier, puis entretien, puis en_attente, puis rejete)
    2. Score decroissant (plus eleve d'abord)
    """
    statut_order = {'retenu': 0, 'entretien': 1, 'en_attente': 2, 'rejete': 3}
    def get_score(candidat):
        try:
            return float(candidat.get('score', 0))
        except (ValueError, TypeError):
            return 0.0
    return sorted(candidats, key=lambda x: (
        statut_order.get(x.get('statut', 'en_attente'), 99),
        -get_score(x)
    ))
def build_ia_prompt(poste, cv_text, lettre_text, attestation_texts_list):
    grille = GRILLE.get(poste, {})
    if not grille:
        return None, None
    regles_interpretation = """
🔴 REGLES D'INTERPRETATION DES POSTES (A RESPECTER ABSOLUMENT) :
1. "Acting Branch Manager" ou "Chef d'Agence" = CHEF D'AGENCE → MANAGEMENT VALIDE (3/3)
2. "Profit Center Manager" = MANAGEMENT D'EQUIPE AVEC P&L → MANAGEMENT VALIDE (3/3)
3. "Chef de service" + "Supervision de l'equipe" = MANAGEMENT D'EQUIPE VALIDE (2-3/3)
4. "Supervision des activites commerciales" = MANAGEMENT VALIDE
5. "Encadrer, diriger et motiver les gestionnaires" = MANAGEMENT D'EQUIPE VALIDE
🔴 REGLES D'INTERPRETATION DES CERTIFICATIONS :
1. "Frankfurt School" = CERTIFICATION BANCAIRE VALABLE (+1)
2. "Moody's" = CERTIFICATION RISQUE VALABLE (+1)
3. "ITB" = CERTIFICATION BANCAIRE VALABLE (+1)
4. "Certificat d'Expert en Financement des PME" = CERTIFICATION VALABLE (+1)
🔴 CAS PARTICULIERS :
1. Chef d'agence (meme interimaire) = EXEMPTE des criteres eliminatoires 4, 5, 6
2. Bac+3 + 10+ ans d'experience = DIPLOME VALIDE (compensation)
3. Les postes de management avec supervision d'equipe sont VALIDES meme sans chiffres
4. Un profil avec 17 ans d'experience bancaire est CONSIDERE comme senior
5. La gestion de portefeuille Local Corporate valide le critere d'experience corporate
6. L'analyse et montage de dossiers de credit valide le critere risque de credit
"""
    system_prompt = f"""Tu es un consultant senior en recrutement bancaire avec 20 ans d'experience en Afrique centrale et de l'Ouest (CEMAC/UEMOA).
POSTE ANALYSE : {poste}
DESCRIPTION : {grille.get('description', 'Poste bancaire')}
Tu dois analyser STRICTEMENT selon la grille officielle fournie. Tu es le SEUL moteur d'analyse.
{regles_interpretation}
REGLES ABSOLUES D'ANALYSE :
1. Tu DOIS comprendre le SENS des phrases et le CONTEXTE METIER.
2. Tu ne JAMAIS inventer des faits qui ne sont PAS dans les documents.
3. Si une information n'est PAS mentionnee, elle est ABSENTE.
4. Les STAGES, BENEVOLATS et FORMATIONS ne comptent PAS comme experience pro.
5. Tu JUSTIFIES chaque evaluation avec des CITATIONS du CV/lettre.
6. Tu utilises le contexte CEMAC/UEMOA (COBAC, BEAC, reglementation locale).
7. Une interpretation RAISONNABLE est attendue : un profil avec 17 ans d'experience bancaire est senior.
8. Les postes de management (Acting Branch Manager, Profit Center Manager, Chef de service) sont consideres comme de l'experience manageriale VALIDE.
🔴 CRITERES ELIMINATOIRES (rejet immediat si non satisfaits) :
{chr(10).join(f"- {c}" for c in grille.get('eliminatoire', []))}
🟠 POINTS A VERIFIER dans le CV :
{chr(10).join(f"- {c}" for c in grille.get('a_verifier', []))}
🟡 SIGNAUX FORTS - candidat prioritaire :
{chr(10).join(f"- {c}" for c in grille.get('signaux_forts', []))}
⚠️ POINTS D'ATTENTION :
{chr(10).join(f"- {c}" for c in grille.get('points_attention', []))}
📊 GRILLE DE SCORING (MAX {grille.get('score_max_total', 14)}) :
{chr(10).join(f"- {k}: {v}/max" for k, v in grille.get('scores_max', {}).items())}
REGLES METIER SPECIFIQUES :
{chr(10).join(f"- {r}" for r in grille.get('regles_metier', []))}
FORMAT DE SORTIE : UNIQUEMENT du JSON valide.
{{
  "flags_eliminatoires": ["liste des criteres eliminatoires non satisfaits"],
  "points_forts": ["liste des points forts avec justifications"],
  "points_vigilance": ["liste des points de vigilance avec justifications"],
  "sous_scores": {{
    {', '.join(f'"{k}": 0-{v}' for k, v in grille.get('scores_max', {}).items())}
  }},
  "synthese_recruteur": "synthese professionnelle pour le recruteur (max 200 mots)",
  "checklist": {{
    "elim_0": true/false,
    "elim_1": true/false,
    "elim_2": true/false,
    "elim_3": true/false,
    "elim_4": true/false,
    "elim_5": true/false,
    "verif_0": true/false,
    "verif_1": true/false,
    "verif_2": true/false,
    "verif_3": true/false,
    "verif_4": true/false,
    "verif_5": true/false,
    "verif_6": true/false
  }},
  "profils_detectes": {{
    "chef_agence": false,
    "gestionnaire_portefeuille": false,
    "local_corporate": false,
    "cross_selling": false
  }},
  "banking_years": 0,
  "diplome_niveau": "Bac+3/Bac+4/Bac+5/Doctorat/Autre"
}}
IMPORTANT : La somme des sous_scores donne le score total. Les flags_eliminatoires doivent etre remplis si des criteres eliminatoires ne sont pas satisfaits. Sois precis et professionnel."""
    user_message = f"""=== DOCUMENTS A ANALYSER ===
CV DU CANDIDAT :
{cv_text[:12000]}
LETTRE DE MOTIVATION :
{lettre_text[:3000] if lettre_text else '(Aucune)'}
ATTESTATIONS :
{''.join(attestation_texts_list)[:3000] if attestation_texts_list else '(Aucune)'}
=== INSTRUCTIONS ===
Analyse strictement selon la grille. Verifie chaque critere eliminatoire.
Le score_total est la SOMME des sous-scores.
Fournis une checklist complete avec true/false pour chaque critere.
Sois rigoureux et professionnel.
INTERPRETE RAISONNABLEMENT les postes de management et les certifications."""
    return system_prompt, user_message
def call_ia_model(client, model, system_prompt, user_message):
    try:
        is_nemotron = "nemotron" in model.lower()
        supports_reasoning = is_nemotron
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            "temperature": 0.1,
            "max_tokens": 4096,
            "extra_headers": {
                "HTTP-Referer": "https://recrutment.onrender.com",
                "X-Title": "RecrutBank CV Analyzer"
            }
        }
        if OPENROUTER_REASONING_ENABLED and supports_reasoning:
            payload["reasoning"] = {"enabled": True}
            logger.info(f"🧠 Reasoning activé pour {model}")
        else:
            payload["response_format"] = {"type": "json_object"}
        response = client.chat.completions.create(**payload)
        if response is None:
            return None
        if not hasattr(response, 'choices') or response.choices is None or len(response.choices) == 0:
            return None
        if response.choices[0].message is None:
            return None
        content = response.choices[0].message.content
        if not content or len(content.strip()) < 10:
            return None
        if supports_reasoning and OPENROUTER_REASONING_ENABLED:
            import re as re_json
            json_match = re_json.search(r'\{[\s\S]*\}', content)
            if json_match:
                try:
                    json.loads(json_match.group(1))
                    return json_match.group(1)
                except:
                    pass
        return content
    except Exception as e:
        logger.warning(f"⚠️ Erreur avec ce modele: {e}")
        return None
def analyze_cv_with_ia_only(cv_text, lettre_text, attestation_texts_list, poste):
    if not IA_ANALYSE_ACTIVE or not ACTIVE_MODELS:
        logger.error("❌ Aucun modele IA disponible")
        return None
    if not cv_text or len(cv_text.strip()) < 50:
        logger.error("❌ CV trop court")
        return None
    if poste not in GRILLE:
        logger.error(f"❌ Poste non trouve dans la grille: {poste}")
        return None
    grille = GRILLE.get(poste, {})
    system_prompt, user_message = build_ia_prompt(poste, cv_text, lettre_text, attestation_texts_list)
    if not system_prompt or not user_message:
        return None
    result_text = None
    used_model = None
    for model_config in ACTIVE_MODELS:
        with _ia_semaphore:
            logger.info(f"🔄 Tentative avec {model_config['name']}...")
            result_text = call_ia_model(
                model_config["client"],
                model_config["model"],
                system_prompt,
                user_message
            )
            if result_text and len(result_text.strip()) > 50:
                used_model = model_config["name"]
                logger.info(f"✅ Analyse reussie avec {used_model}")
                break
            else:
                logger.warning(f"⚠️ {model_config['name']} n'a pas retourne de reponse valide")
    if not result_text or len(result_text.strip()) < 50:
        logger.error("❌ Tous les modeles IA ont echoue")
        return None
    logger.info(f"✅ Analyse IA terminee: {len(result_text)} caracteres avec {used_model}")
    analyse = parse_json_robust(result_text)
    if analyse is None:
        logger.error("❌ Echec de l'extraction JSON")
        return None
    sous_scores = analyse.get('sous_scores', {})
    max_scores = grille.get('scores_max', {})
    score_max_total = grille.get('score_max_total', 14)
    score_total = 0
    for key, max_val in max_scores.items():
        val = sous_scores.get(key, 0)
        if val > max_val:
            val = max_val
        if val < 0:
            val = 0
        sous_scores[key] = val
        score_total += val
    score_total = min(score_max_total, max(0, score_total))
    flags_elim = analyse.get('flags_eliminatoires', [])
    if not isinstance(flags_elim, list):
        flags_elim = []
    veritable_eliminatoire = grille.get('eliminatoire', [])
    flags_elim_corriges = []
    for flag in flags_elim:
        for critere in veritable_eliminatoire:
            if critere.lower() in flag.lower() or flag.lower() in critere.lower():
                flags_elim_corriges.append(flag)
                break
    flags_elim = flags_elim_corriges
    checklist = analyse.get('checklist', {})
    if not checklist or len(checklist) == 0:
        keys_list = list(max_scores.keys())
        checklist = {
            "elim_0": len(flags_elim) == 0,
            "elim_1": len(flags_elim) == 0,
            "elim_2": len(flags_elim) == 0,
            "elim_3": len(flags_elim) == 0,
            "elim_4": len(flags_elim) == 0,
            "elim_5": len(flags_elim) == 0,
            "verif_0": sous_scores.get(keys_list[0], 0) >= 2 if keys_list else False,
            "verif_1": sous_scores.get(keys_list[1], 0) >= 2 if len(keys_list) > 1 else False,
            "verif_2": sous_scores.get(keys_list[2], 0) >= 1 if len(keys_list) > 2 else False,
            "verif_3": sous_scores.get(keys_list[3], 0) >= 1 if len(keys_list) > 3 else False,
            "verif_4": sous_scores.get(keys_list[4], 0) >= 1 if len(keys_list) > 4 else False,
            "verif_5": sous_scores.get(keys_list[5], 0) >= 1 if len(keys_list) > 5 else False,
            "verif_6": sous_scores.get(keys_list[6], 0) >= 1 if len(keys_list) > 6 else False
        }
    decision = get_recommandation_from_score(score_total, flags_elim, score_max_total)
    statut = get_statut_from_decision(decision, flags_elim)
    profils = analyse.get('profils_detectes', {})
    banking_years = analyse.get('banking_years', 0)
    result = {
        'poste': poste,
        'score': score_total,
        'score_max': score_max_total,
        'decision': decision,
        'statut': statut,
        'flags_eliminatoires': flags_elim,
        'sous_scores': sous_scores,
        'checklist': checklist,
        'points_forts': analyse.get('points_forts', []),
        'points_vigilance': analyse.get('points_vigilance', []),
        'synthese': analyse.get('synthese_recruteur', ''),
        'score_breakdown': {
            'score_final': score_total,
            'score_max': score_max_total,
            'decision': decision,
            'statut': statut,
            'sous_scores': sous_scores,
            'chef_agence_detecte': profils.get('chef_agence', False),
            'gestionnaire_portefeuille_detecte': profils.get('gestionnaire_portefeuille', False),
            'local_corporate_detecte': profils.get('local_corporate', False),
            'cross_selling_detecte': profils.get('cross_selling', False),
            'banking_years_detecte': banking_years,
            'nb_eliminatoires': len(flags_elim),
            'eliminatoires_passes': len(flags_elim) == 0,
            'grille_version': 'v13.1-ia-multi-fallback',
            'moteur': 'IA_100%',
            'model_used': used_model
        },
        'analyse_details': {
            'moteur': 'IA_100%',
            'provider': used_model,
            'model_used': used_model,
            'sous_scores': sous_scores,
            'flags_eliminatoires': flags_elim,
            'points_forts': analyse.get('points_forts', []),
            'points_vigilance': analyse.get('points_vigilance', []),
            'synthese_recruteur': analyse.get('synthese_recruteur', ''),
            'profils_detectes': profils,
            'banking_years': banking_years,
            'diplome_niveau': analyse.get('diplome_niveau', 'Non specifie')
        }
    }
    return result
def run_analysis_for_candidat(token, cv_filename, lettre_filename, attestation_filenames, poste, force=False):
    try:
        if not force and not is_poste_actif(poste):
            logger.info(f"Analyse ignoree pour {token} — poste cloture : {poste}")
            if supabase:
                supabase.table('candidats').update({
                    "analyse_status": "skipped_closed_post",
                    "analyse_auto_date": datetime.datetime.now().isoformat()
                }).eq('token', token).execute()
            return
        if isinstance(attestation_filenames, str):
            try:
                attestation_filenames = json.loads(attestation_filenames) if attestation_filenames else []
            except Exception:
                attestation_filenames = [attestation_filenames] if attestation_filenames else []
        cv_text = ""
        if cv_filename:
            cv_bytes = download_file_from_supabase_robust(cv_filename)
            if cv_bytes:
                cv_text = extract_text_robust_from_bytes(cv_bytes, cv_filename)
                if len(cv_text) > MAX_TEXT_SIZE:
                    cv_text = cv_text[:MAX_TEXT_SIZE]
                logger.info(f"CV extrait pour {token}: {len(cv_text)} caracteres")
        if not cv_text or len(cv_text.strip()) < 30:
            logger.warning(f"CV manquant ou vide pour {token}")
            if supabase:
                supabase.table('candidats').update({
                    "score": "0",
                    "decision": "Rejet - CV manquant",
                    "statut": "rejete",
                    "analyse_status": "error"
                }).eq('token', token).execute()
            return
        lm_text = ""
        if lettre_filename:
            lm_bytes = download_file_from_supabase_robust(lettre_filename)
            if lm_bytes:
                lm_text = extract_text_robust_from_bytes(lm_bytes, lettre_filename)
                if len(lm_text) > MAX_TEXT_SIZE:
                    lm_text = lm_text[:MAX_TEXT_SIZE]
                logger.info(f"Lettre extraite pour {token}: {len(lm_text)} caracteres")
        att_texts = []
        for fn in (attestation_filenames or []):
            if fn:
                att_bytes = download_file_from_supabase_robust(fn)
                if att_bytes:
                    t = extract_text_robust_from_bytes(att_bytes, fn)
                    if t and len(t) > MAX_TEXT_SIZE:
                        t = t[:MAX_TEXT_SIZE]
                    if t:
                        att_texts.append(t)
        if not IA_ANALYSE_ACTIVE or not ACTIVE_MODELS:
            logger.error("❌ Aucun modele IA disponible - impossible d'analyser")
            if supabase:
                supabase.table('candidats').update({
                    "score": "0",
                    "decision": "Erreur - IA non disponible",
                    "statut": "rejete",
                    "analyse_status": "error"
                }).eq('token', token).execute()
            return
        result = analyze_cv_with_ia_only(cv_text, lm_text, att_texts, poste)
        if not result:
            logger.error(f"❌ Analyse IA echouee pour {token} (tous les modeles)")
            if supabase:
                supabase.table('candidats').update({
                    "analyse_status": "error",
                    "analyse_error": "L'analyse IA a echoue avec tous les modeles"
                }).eq('token', token).execute()
            return
        score = result.get('score', 0)
        score_max = result.get('score_max', 14)
        decision = result.get('decision', 'Rejet')
        statut = result.get('statut', 'rejete')
        flags_elim = result.get('flags_eliminatoires', [])
        sous_scores = result.get('sous_scores', {})
        score_breakdown = result.get('score_breakdown', {})
        checklist = result.get('checklist', {})
        points_forts = result.get('points_forts', [])
        points_vigilance = result.get('points_vigilance', [])
        synthese = result.get('synthese', '')
        analyse_details = result.get('analyse_details', {})
        if supabase:
            update_data = {
                "score": str(score),
                "decision": decision,
                "statut": statut,
                "analyse_status": "completed",
                "analyse_auto_date": datetime.datetime.now().isoformat(),
                "flags_eliminatoires": json.dumps(flags_elim, ensure_ascii=False),
                "score_breakdown": json.dumps(score_breakdown, ensure_ascii=False),
                "analyse_details": json.dumps(analyse_details, ensure_ascii=False),
                "checklist": json.dumps(checklist, ensure_ascii=False)
            }
            supabase.table('candidats').update(update_data).eq('token', token).execute()
        model_used = analyse_details.get('model_used', 'Inconnu')
        logger.info(f"[{decision}] Score {token}: {score}/{score_max} → statut: {statut} (modele: {model_used}, flags: {len(flags_elim)})")
    except Exception as e:
        import traceback
        traceback.print_exc()
        logger.error(f"Erreur analyse {token}: {str(e)}")
        if supabase:
            try:
                supabase.table('candidats').update({
                    "analyse_status": "error",
                    "analyse_error": str(e)
                }).eq('token', token).execute()
            except:
                pass
def generate_excel_report_enhanced(candidats, poste_filter=""):
    if not OPENPYXL_AVAILABLE:
        return None
    try:
        sorted_candidats = sort_candidats(candidats)
        wb = Workbook()
        ws = wb.active
        ws.title = "Candidats"
        header_font = Font(bold=True, size=12, color="FFFFFF")
        header_fill = PatternFill(start_color="1a3a5c", end_color="1a3a5c", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell_alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        center_alignment = Alignment(horizontal="center", vertical="center")
        number_alignment = Alignment(horizontal="center", vertical="center")
        thin_border = Border(
            left=Side(style='thin', color='000000'),
            right=Side(style='thin', color='000000'),
            top=Side(style='thin', color='000000'),
            bottom=Side(style='thin', color='000000')
        )
        headers = [
            "Rang", "N° Dossier", "Nom", "Prénom", "Email", "Téléphone",
            "Poste", "Statut", "Score", "Décision",
            "Points Forts", "Points de Vigilance", "Synthèse"
        ]
        col_widths = [5, 12, 18, 18, 25, 15, 30, 12, 8, 20, 35, 35, 40]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
            ws.column_dimensions[get_column_letter(col)].width = col_widths[col-1]
        row_idx = 2
        rank = 1
        for c in sorted_candidats:
            score_breakdown = c.get('score_breakdown_parsed', {})
            analyse_details = c.get('analyse_details_parsed', {})
            points_forts = analyse_details.get('points_forts', []) or score_breakdown.get('points_forts', [])
            points_vigilance = analyse_details.get('points_vigilance', []) or score_breakdown.get('points_vigilance', [])
            synthese = analyse_details.get('synthese_recruteur', '') or c.get('synthese', '')
            statut = c.get('statut', 'en_attente')
            try:
                score = float(c.get('score', 0))
            except (ValueError, TypeError):
                score = 0.0
            ws.cell(row=row_idx, column=1, value=rank).alignment = number_alignment
            ws.cell(row=row_idx, column=2, value=c.get('numero_dossier', '')).alignment = center_alignment
            ws.cell(row=row_idx, column=3, value=c.get('nom', '')).alignment = cell_alignment
            ws.cell(row=row_idx, column=4, value=c.get('prenom', '')).alignment = cell_alignment
            ws.cell(row=row_idx, column=5, value=c.get('email', '')).alignment = cell_alignment
            ws.cell(row=row_idx, column=6, value=c.get('telephone', '')).alignment = cell_alignment
            ws.cell(row=row_idx, column=7, value=c.get('poste', '')).alignment = cell_alignment
            ws.cell(row=row_idx, column=8, value=statut).alignment = center_alignment
            ws.cell(row=row_idx, column=9, value=score).alignment = center_alignment
            ws.cell(row=row_idx, column=10, value=c.get('decision', '')).alignment = cell_alignment
            ws.cell(row=row_idx, column=11, value=", ".join(points_forts[:4]) if points_forts else '').alignment = cell_alignment
            ws.cell(row=row_idx, column=12, value=", ".join(points_vigilance[:4]) if points_vigilance else '').alignment = cell_alignment
            ws.cell(row=row_idx, column=13, value=synthese[:300] if synthese else '').alignment = cell_alignment
            for col in range(1, len(headers) + 1):
                ws.cell(row=row_idx, column=col).border = thin_border
            if row_idx % 2 == 0:
                for col in range(1, len(headers) + 1):
                    ws.cell(row=row_idx, column=col).fill = PatternFill(start_color="f8f9fa", end_color="f8f9fa", fill_type="solid")
            row_idx += 1
            rank += 1
        total_row = row_idx
        if len(sorted_candidats) > 0:
            total_cell = ws.cell(row=total_row, column=1, value=f"Total: {len(sorted_candidats)} candidats")
            total_cell.font = Font(bold=True, size=11)
            ws.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=3)
            total_cell.alignment = Alignment(horizontal="left", vertical="center")
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf
    except Exception as e:
        logger.error(f"Erreur generation Excel: {e}")
        return None
def generate_pdf_report_enhanced(candidats, poste_filter=""):
    if not REPORTLAB_AVAILABLE:
        return None
    try:
        sorted_candidats = sort_candidats(candidats)
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=1.2*cm, leftMargin=1.2*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            textColor=colors.HexColor('#1a3a5c'),
            alignment=TA_CENTER,
            spaceAfter=6
        )
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#666666'),
            alignment=TA_CENTER,
            spaceAfter=12
        )
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.whitesmoke,
            alignment=TA_CENTER,
            backColor=colors.HexColor('#1a3a5c')
        )
        cell_style = ParagraphStyle(
            'CustomCell',
            parent=styles['Normal'],
            fontSize=8,
            alignment=TA_LEFT,
            leading=10
        )
        cell_center = ParagraphStyle(
            'CustomCellCenter',
            parent=styles['Normal'],
            fontSize=8,
            alignment=TA_CENTER,
            leading=10
        )
        story = []
        story.append(Paragraph("RAPPORT DE CANDIDATURES", title_style))
        story.append(Paragraph(f"Poste: {poste_filter if poste_filter else 'Tous les postes'} | Genere le {datetime.datetime.now().strftime('%d/%m/%Y a %H:%M')}", subtitle_style))
        story.append(Spacer(1, 0.3*cm))
        if not sorted_candidats:
            story.append(Paragraph("Aucun candidat trouve.", styles['Normal']))
        else:
            data = [
                [
                    Paragraph("Rang", header_style),
                    Paragraph("N° Dossier", header_style),
                    Paragraph("Nom", header_style),
                    Paragraph("Prenom", header_style),
                    Paragraph("Email", header_style),
                    Paragraph("Poste", header_style),
                    Paragraph("Statut", header_style),
                    Paragraph("Score", header_style)
                ]
            ]
            rank = 1
            for c in sorted_candidats[:100]:
                statut = c.get('statut', 'en_attente')
                try:
                    score = float(c.get('score', 0))
                except (ValueError, TypeError):
                    score = 0.0
                if statut == 'retenu':
                    statut_text = f'<font color="#16a34a">✅ Retenu</font>'
                elif statut == 'entretien':
                    statut_text = f'<font color="#d97706">📅 Entretien</font>'
                elif statut == 'rejete':
                    statut_text = f'<font color="#dc2626">❌ Rejete</font>'
                else:
                    statut_text = f'<font color="#d97706">⏳ En attente</font>'
                data.append([
                    Paragraph(str(rank), cell_center),
                    Paragraph(c.get('numero_dossier', ''), cell_center),
                    Paragraph(c.get('nom', ''), cell_style),
                    Paragraph(c.get('prenom', ''), cell_style),
                    Paragraph(c.get('email', ''), cell_style),
                    Paragraph(c.get('poste', '')[:30], cell_style),
                    Paragraph(statut_text, cell_center),
                    Paragraph(str(score), cell_center)
                ])
                rank += 1
            col_widths = [1.2*cm, 2.5*cm, 3.5*cm, 3.5*cm, 4.5*cm, 4.5*cm, 2.8*cm, 1.5*cm]
            table = Table(data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a3a5c')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                ('TOPPADDING', (0, 0), (-1, 0), 6),
                ('GRID', (0, 0), (-1, 0), 0.5, colors.HexColor('#333333')),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                ('ALIGN', (0, 1), (0, -1), 'CENTER'),
                ('ALIGN', (1, 1), (1, -1), 'CENTER'),
                ('ALIGN', (6, 1), (7, -1), 'CENTER'),
                ('GRID', (0, 1), (-1, -1), 0.3, colors.HexColor('#CCCCCC')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#FFFFFF'), colors.HexColor('#F8F9FA')]),
            ]))
            story.append(table)
            story.append(Spacer(1, 0.5*cm))
            story.append(Paragraph(f"Total: {len(sorted_candidats)} candidat(s)", styles['Normal']))
        doc.build(story)
        buf.seek(0)
        return buf
    except Exception as e:
        logger.error(f"Erreur generation PDF: {e}")
        return None
def generate_csv_report_enhanced(candidats, poste_filter=""):
    sorted_candidats = sort_candidats(candidats)
    output = io.StringIO()
    writer = csv.writer(output, delimiter=';')
    headers = ["Rang", "N° Dossier", "Nom", "Prénom", "Email", "Téléphone", "Poste", "Statut", "Score", "Décision"]
    writer.writerow(headers)
    rank = 1
    for c in sorted_candidats:
        writer.writerow([
            rank,
            c.get('numero_dossier', ''),
            c.get('nom', ''),
            c.get('prenom', ''),
            c.get('email', ''),
            c.get('telephone', ''),
            c.get('poste', ''),
            c.get('statut', 'en_attente'),
            c.get('score', 0),
            c.get('decision', '')
        ])
        rank += 1
    return output.getvalue()
def generate_word_report_enhanced(candidats, poste_filter=""):
    if not DOCX_AVAILABLE:
        return None
    try:
        sorted_candidats = sort_candidats(candidats)
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        doc = DocxDocument()
        def set_cell_border(cell, **kwargs):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for edge in ['top', 'left', 'bottom', 'right']:
                tag = f'w:{edge}'
                if edge in kwargs:
                    border = OxmlElement(tag)
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'CCCCCC')
                    tcPr.append(border)
        title = doc.add_heading(f"RAPPORT DE CANDIDATURES", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Poste: {poste_filter if poste_filter else 'Tous les postes'} | Genere le {datetime.datetime.now().strftime('%d/%m/%Y a %H:%M')}")
        doc.add_paragraph()
        if not sorted_candidats:
            doc.add_paragraph("Aucun candidat trouve.")
        else:
            doc.add_heading("Liste des candidats", level=1)
            table = doc.add_table(rows=1, cols=10)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            header_cells = table.rows[0].cells
            headers = ["Rang", "N° Dossier", "Nom", "Prénom", "Email", "Poste", "Statut", "Score", "Décision"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].font.bold = True
                header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
                header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'solid')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), '1a3a5c')
                header_cells[i]._tc.get_or_add_tcPr().append(shading)
            rank = 1
            for c in sorted_candidats[:200]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(rank)
                row_cells[1].text = c.get('numero_dossier', '')
                row_cells[2].text = c.get('nom', '')
                row_cells[3].text = c.get('prenom', '')
                row_cells[4].text = c.get('email', '')
                row_cells[5].text = c.get('poste', '')[:40]
                row_cells[6].text = c.get('statut', 'en_attente')
                row_cells[7].text = str(c.get('score', 0))
                row_cells[8].text = c.get('decision', '')[:30]
                for idx in [0, 1, 6, 7]:
                    row_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                rank += 1
            for i, width in enumerate([1.2, 2.5, 3.5, 3.5, 4.5, 4.5, 2.8, 1.5, 4.0]):
                table.columns[i].width = Cm(width)
            doc.add_paragraph()
            doc.add_paragraph(f"Total: {len(sorted_candidats)} candidat(s)")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    except Exception as e:
        logger.error(f"Erreur generation Word: {e}")
        return None
@app.route('/api/postes', methods=['GET'])
def get_postes():
    return jsonify({"postes": POSTES, "postes_actifs": POSTES_ACTIFS, "postes_clotures": POSTES_CLOTURES}), 200
@app.route('/api/postes/actifs', methods=['GET'])
def get_postes_actifs():
    return jsonify(POSTES_ACTIFS), 200
@app.route('/api/grille/<poste>', methods=['GET'])
def get_grille(poste):
    g = GRILLE.get(poste)
    if not g:
        return jsonify({'error': 'Poste inconnu', 'postes_disponibles': list(GRILLE.keys())}), 404
    return jsonify(g), 200
@app.route('/api/auth/login', methods=['POST', 'OPTIONS'])
def login():
    if request.method == 'OPTIONS':
        return '', 204
    data = request.get_json(silent=True)
    if not data:
        return jsonify({'error': 'JSON manquant'}), 400
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')
    hashed_password = hash_pwd(password)
    if supabase:
        try:
            response = supabase.table('recruteurs').select('*').eq('email', email).execute()
            if response.data and len(response.data) > 0:
                recruteur = response.data[0]
                if recruteur.get('password') == hashed_password:
                    access_token = create_access_token(identity=str(recruteur['id']))
                    return jsonify({
                        'token': access_token,
                        'nom': recruteur.get('nom', 'Recruteur'),
                        'email': recruteur.get('email', email)
                    }), 200
        except Exception as e:
            logger.error(f"Erreur login: {e}")
    return jsonify({'error': 'Identifiants incorrects'}), 401
@app.route('/api/candidats/postuler', methods=['POST'])
def postuler():
    try:
        nom = (request.form.get('nom') or '').strip()
        prenom = (request.form.get('prenom') or '').strip()
        email = (request.form.get('email') or '').strip().lower()
        telephone = (request.form.get('telephone') or '').strip()
        poste = (request.form.get('poste') or '').strip()
        if not nom or not prenom or not email or poste not in POSTES:
            return jsonify({'error': 'Champs obligatoires manquants ou poste invalide'}), 400
        if supabase:
            existing = supabase.table('candidats').select('*').eq('email', email).eq('poste', poste).execute()
            if existing.data and len(existing.data) > 0:
                return jsonify({'error': f'Vous avez deja soumis une candidature pour le poste "{poste}".'}), 409
            all_candidats = supabase.table('candidats').select('numero_dossier').eq('poste', poste).execute()
            max_num = 0
            for c in all_candidats.data:
                existing_num = c.get('numero_dossier', '')
                if existing_num:
                    try:
                        num_val = int(existing_num)
                        if num_val > max_num:
                            max_num = num_val
                    except ValueError:
                        pass
            new_num = max_num + 1
            numero_dossier = str(new_num)
        def save_file_to_supabase(field, suffix):
            f = request.files.get(field)
            if f and f.filename and allowed_file(f.filename):
                ext = f.filename.rsplit('.', 1)[-1].lower()
                blob_name = f"{uuid.uuid4().hex}_{suffix}.{ext}"
                result = upload_file_to_supabase(f, blob_name, f.content_type)
                return result if result else ''
            return ''
        cv_filename = save_file_to_supabase('cv', 'cv')
        if request.files.get('cv') and not cv_filename:
            return jsonify({'error': "Echec de l'envoi du CV, merci de reessayer."}), 500
        lettre_filename = save_file_to_supabase('lettre', 'lettre')
        if request.files.get('lettre') and not lettre_filename:
            return jsonify({'error': "Echec de l'envoi de la lettre de motivation, merci de reessayer."}), 500
        att_filenames = []
        for f in request.files.getlist('attestation'):
            if f and f.filename and allowed_file(f.filename):
                ext = f.filename.rsplit('.', 1)[-1].lower()
                blob_name = f"{uuid.uuid4().hex}_attestation.{ext}"
                result = upload_file_to_supabase(f, blob_name, f.content_type)
                if result:
                    att_filenames.append(blob_name)
        token = uuid.uuid4().hex
        supabase.table('candidats').insert({
            "token": token,
            "nom": nom,
            "prenom": prenom,
            "email": email,
            "telephone": telephone,
            "poste": poste,
            "numero_dossier": numero_dossier,
            "cv_filename": cv_filename,
            "lettre_filename": lettre_filename,
            "attestation_filenames": json.dumps(att_filenames, ensure_ascii=False),
            "statut": "en_attente",
            "note": "",
            "score": "0",
            "checklist": "",
            "flags_eliminatoires": "",
            "score_breakdown": "",
            "analyse_status": "pending",
            "date_candidature": datetime.datetime.now().isoformat()
        }).execute()
        if is_poste_actif(poste):
            threading.Thread(target=run_analysis_for_candidat, args=(token, cv_filename, lettre_filename, att_filenames, poste, False), daemon=True).start()
            analyse_msg = f'Analyse automatique en cours avec {len(ACTIVE_MODELS)} modele(s) IA'
        else:
            analyse_msg = 'Poste cloture — candidature enregistree sans analyse'
            supabase.table('candidats').update({
                "analyse_status": "closed_post_no_analysis",
                "analyse_auto_date": datetime.datetime.now().isoformat()
            }).eq('token', token).execute()
        nom_complet = f"{prenom} {nom}".strip()
        sujet_confirmation = f"Confirmation de candidature – {poste}"
        corps_confirmation = f"""Bonjour {nom_complet},
Nous accusons reception de votre candidature pour le poste de {poste}.
Votre dossier sera analyse par notre systeme d'intelligence artificielle dans les plus brefs delais.
Vous recevrez une notification lorsque l'analyse sera terminee.
Sans reponse de notre part sous deux (2) semaines, veuillez considerer que votre candidature n'a pas ete retenue.
Pour toute information : contact@cdotchad.com.
Cordialement,
L'equipe RecrutBank"""
        threading.Thread(target=send_email, args=(email, sujet_confirmation, corps_confirmation), daemon=True).start()
        return jsonify({
            'message': 'Candidature soumise avec succes',
            'token': token,
            'numero_dossier': numero_dossier,
            'analyse': analyse_msg,
            'poste_statut': 'actif' if is_poste_actif(poste) else 'cloture',
            'ia_engine': f"{len(ACTIVE_MODELS)} modeles IA"
        }), 201
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
@app.route('/api/candidats/statut/<token>', methods=['GET'])
def get_statut(token):
    if supabase:
        response = supabase.table('candidats').select('*').eq('token', token).execute()
        if response.data and len(response.data) > 0:
            data = response.data[0]
            hidden = {'cv_filename', 'lettre_filename', 'attestation_filenames', 'checklist', 'flags_eliminatoires', 'analyse_details', 'score_breakdown'}
            return jsonify({k: v for k, v in data.items() if k not in hidden}), 200
    return jsonify({'error': 'Candidature introuvable'}), 404
@app.route('/api/recruteur/stats', methods=['GET'])
@jwt_required()
def get_stats():
    if not supabase:
        return jsonify({'error': 'Supabase non configure'}), 500
    response = supabase.table('candidats').select('*').execute()
    keys = response.data if response.data else []
    stats = {"total": len(keys), "en_attente": 0, "retenu": 0, "rejete": 0, "entretien": 0, "by_poste": []}
    counts = {}
    for c in keys:
        statut = c.get('statut', 'en_attente')
        if statut in stats:
            stats[statut] += 1
        else:
            stats['en_attente'] += 1
        p = c.get('poste', 'Inconnu')
        counts[p] = counts.get(p, 0) + 1
    stats['by_poste'] = [{'poste': p, 'n': n} for p, n in sorted(counts.items(), key=lambda x: -x[1])]
    return jsonify(stats), 200
@app.route('/api/recruteur/postes/stats', methods=['GET'])
@jwt_required()
def get_postes_stats():
    if not supabase:
        return jsonify({'error': 'Supabase non configure'}), 500
    response = supabase.table('candidats').select('*').execute()
    keys = response.data if response.data else []
    actifs_count = 0
    clotures_count = 0
    par_poste_actif = {}
    par_poste_cloture = {}
    for c in keys:
        poste = c.get('poste', '')
        if poste in POSTES_ACTIFS:
            actifs_count += 1
            par_poste_actif[poste] = par_poste_actif.get(poste, 0) + 1
        else:
            clotures_count += 1
            par_poste_cloture[poste] = par_poste_cloture.get(poste, 0) + 1
    return jsonify({
        'total': len(keys),
        'postes_actifs': {
            'count': actifs_count,
            'liste': POSTES_ACTIFS,
            'par_poste': par_poste_actif,
            'eligible_reanalyse': True
        },
        'postes_clotures': {
            'count': clotures_count,
            'liste': POSTES_CLOTURES,
            'par_poste': par_poste_cloture,
            'eligible_reanalyse': False
        }
    }), 200
@app.route('/api/recruteur/candidats', methods=['GET'])
@jwt_required()
def list_candidats():
    poste_filter = request.args.get('poste', '')
    statut_filter = request.args.get('statut', '')
    search = request.args.get('search', '').lower()
    min_score = request.args.get('min_score', type=int)
    if not supabase:
        return jsonify({'error': 'Supabase non configure'}), 500
    response = supabase.table('candidats').select('*').execute()
    all_candidats = response.data if response.data else []
    result = []
    for c in all_candidats:
        c['id'] = c.get('token', '')
        if poste_filter and c.get('poste') != poste_filter:
            continue
        if statut_filter:
            statut = c.get('statut', 'en_attente')
            if statut != statut_filter:
                continue
        if min_score is not None and int(c.get('score', 0)) < min_score:
            continue
        if search:
            hay = (f"{c.get('nom','')} {c.get('prenom','')} {c.get('email','')} {c.get('poste','')} {c.get('numero_dossier','')}").lower()
            if search not in hay:
                continue
        for field in ['score_breakdown', 'flags_eliminatoires', 'analyse_details']:
            if c.get(field):
                try:
                    c[f'{field}_parsed'] = json.loads(c[field])
                except Exception:
                    pass
        result.append(c)
    result = sort_candidats(result)
    return jsonify(result), 200
@app.route('/api/recruteur/candidats/<token>', methods=['GET'])
@jwt_required()
def get_candidat_detail(token):
    if not supabase:
        return jsonify({'error': 'Supabase non configure'}), 500
    response = supabase.table('candidats').select('*').eq('token', token).execute()
    if not response.data or len(response.data) == 0:
        return jsonify({'error': 'Candidat introuvable'}), 404
    data = response.data[0]
    data['id'] = token
    if data.get('attestation_filenames'):
        try:
            data['attestation_filenames_parsed'] = json.loads(data['attestation_filenames'])
        except Exception:
            data['attestation_filenames_parsed'] = []
    for field in ['checklist', 'flags_eliminatoires', 'analyse_details', 'score_breakdown']:
        if data.get(field):
            try:
                data[f'{field}_parsed'] = json.loads(data[field])
            except Exception:
                pass
    return jsonify(data), 200
@app.route('/api/recruteur/candidats/<token>/statut', methods=['PUT'])
@jwt_required()
def update_candidat(token):
    if not supabase:
        return jsonify({'error': 'Supabase non configure'}), 500
    response = supabase.table('candidats').select('*').eq('token', token).execute()
    if not response.data or len(response.data) == 0:
        return jsonify({'error': 'Candidat introuvable'}), 404
    data = request.get_json(silent=True) or {}
    statut = data.get('statut', 'en_attente')
    note = data.get('note', '')
    if statut not in ('en_attente', 'retenu', 'rejete', 'entretien'):
        return jsonify({'error': 'Statut invalide'}), 400
    update_data = {
        "statut": statut,
        "note": note,
        "decision_date": datetime.datetime.now().isoformat(),
        "decided_by": get_jwt_identity()
    }
    if statut == "rejete":
        update_data["decision"] = "Rejet - Decision du recruteur"
    elif statut == "retenu":
        update_data["decision"] = "Retenu - Decision du recruteur"
    elif statut == "entretien":
        update_data["decision"] = "Entretien - Decision du recruteur"
    supabase.table('candidats').update(update_data).eq('token', token).execute()
    return jsonify({'message': 'Mis a jour avec succes', 'statut': statut}), 200
@app.route('/api/recruteur/candidats/<token>/analyze', methods=['POST'])
@jwt_required()
def trigger_analyze(token):
    if not supabase:
        return jsonify({'error': 'Supabase non configure'}), 500
    response = supabase.table('candidats').select('*').eq('token', token).execute()
    if not response.data or len(response.data) == 0:
        return jsonify({'error': 'Candidat introuvable'}), 404
    data = response.data[0]
    cv_fn = data.get('cv_filename')
    lm_fn = data.get('lettre_filename')
    att_raw = data.get('attestation_filenames', '[]')
    poste = data.get('poste')
    if not cv_fn:
        return jsonify({'error': 'CV manquant pour analyse'}), 400
    force = request.args.get('force', '0') == '1'
    if not force and not is_poste_actif(poste):
        return jsonify({
            'error': f'Le poste "{poste}" est cloture. Utilisez ?force=1 pour forcer l\'analyse.',
            'poste': poste,
            'statut': 'cloture'
        }), 403
    supabase.table('candidats').update({
        "analyse_status": "pending",
        "analyse_manual_trigger": datetime.datetime.now().isoformat()
    }).eq('token', token).execute()
    threading.Thread(target=run_analysis_for_candidat, args=(token, cv_fn, lm_fn, att_raw, poste, force), daemon=True).start()
    return jsonify({
        'message': f'Analyse relancee avec {len(ACTIVE_MODELS)} modele(s) IA',
        'token': token,
        'ia_engine': f"{len(ACTIVE_MODELS)} modeles"
    }), 202
@app.route('/api/recruteur/reanalyze-status', methods=['GET'])
@jwt_required()
def get_reanalyze_status():
    try:
        if not supabase:
            return jsonify({'error': 'Supabase non configure'}), 500
        response = supabase.table('candidats').select('token, poste, analyse_status, analyse_auto_date').execute()
        keys = response.data if response.data else []
        active_candidates = [d for d in keys if d.get('poste') in POSTES_ACTIFS]
        total = len(active_candidates)
        status_counts = {'pending': 0, 'reanalyzing': 0, 'completed': 0, 'error': 0, 'skipped_closed_post': 0, 'closed_post_no_analysis': 0, 'reanalyzing_auto': 0}
        in_progress = False
        for data in active_candidates:
            status = data.get('analyse_status', 'pending')
            if status in status_counts:
                status_counts[status] += 1
            if status in ('reanalyzing', 'pending'):
                in_progress = True
        processed = status_counts.get('completed', 0) + status_counts.get('error', 0)
        return jsonify({
            'total': total,
            'processed': processed,
            'in_progress': in_progress,
            'status_counts': status_counts,
            'postes_concernes': POSTES_ACTIFS,
            'timestamp': datetime.datetime.now().isoformat()
        }), 200
    except Exception as e:
        logger.error(f"Erreur reanalyze-status: {e}")
        return jsonify({'error': str(e)}), 500
@app.route('/api/recruteur/reanalyze-one/<token>', methods=['POST'])
@jwt_required()
def reanalyze_one_candidate(token):
    try:
        if not supabase:
            return jsonify({'error': 'Supabase non configure'}), 500
        response = supabase.table('candidats').select('*').eq('token', token).execute()
        if not response.data or len(response.data) == 0:
            return jsonify({'error': 'Candidat introuvable'}), 404
        data = response.data[0]
        poste = data.get('poste')
        if not is_poste_actif(poste):
            return jsonify({
                'error': f'Le poste "{poste}" est cloture. Reanalyse desactivee.',
                'poste': poste,
                'statut': 'cloture'
            }), 403
        cv_fn = data.get('cv_filename')
        if not cv_fn:
            return jsonify({'error': 'CV manquant pour analyse'}), 400
        lm_fn = data.get('lettre_filename')
        att_raw = data.get('attestation_filenames', '[]')
        supabase.table('candidats').update({
            "analyse_status": "reanalyzing",
            "reanalyze_trigger": datetime.datetime.now().isoformat(),
            "reanalyze_reason": "Reanalyse manuelle (un seul candidat)"
        }).eq('token', token).execute()
        threading.Thread(target=run_analysis_for_candidat, args=(token, cv_fn, lm_fn, att_raw, poste, True), daemon=True).start()
        return jsonify({'message': 'Reanalyse lancee (multi-modeles IA)', 'token': token, 'poste': poste}), 202
    except Exception as e:
        logger.error(f"Erreur reanalyze-one: {e}")
        return jsonify({'error': str(e)}), 500
@app.route('/api/recruteur/reanalyze-all', methods=['POST'])
@jwt_required()
def reanalyze_all_candidates():
    try:
        if not supabase:
            return jsonify({'error': 'Supabase non configure'}), 500
        response = supabase.table('candidats').select('*').execute()
        keys = response.data if response.data else []
        if not keys:
            return jsonify({'message': 'Aucune candidature a reanalyser'}), 200
        candidates_to_reanalyze = [data for data in keys if data.get('poste') in POSTES_ACTIFS and data.get('cv_filename')]
        candidates_skipped = len(keys) - len(candidates_to_reanalyze)
        if not candidates_to_reanalyze:
            return jsonify({
                'message': 'Aucun candidat sur poste actif avec CV a reanalyser',
                'skipped_closed_posts': candidates_skipped
            }), 200
        now_iso = datetime.datetime.now().isoformat()
        for c in candidates_to_reanalyze:
            try:
                supabase.table('candidats').update({
                    "analyse_status": "reanalyzing",
                    "reanalyze_trigger": now_iso,
                    "reanalyze_reason": "Reanalyse parallellisee (postes actifs)"
                }).eq('token', c.get('token')).execute()
            except Exception:
                pass
        def analyze_one(data):
            try:
                token = data.get('token')
                cv_fn = data.get('cv_filename')
                lm_fn = data.get('lettre_filename')
                att_raw = data.get('attestation_filenames', '[]')
                poste = data.get('poste')
                if not cv_fn:
                    return (token, False, "CV manquant")
                run_analysis_for_candidat(token, cv_fn, lm_fn, att_raw, poste, False)
                return (token, True, "OK")
            except Exception as e:
                return (data.get('token'), False, str(e))
        MAX_WORKERS = min(8, len(candidates_to_reanalyze))
        logger.info(f"Reanalyse parallele : {len(candidates_to_reanalyze)} candidats, {MAX_WORKERS} workers")
        start_time = time.time()
        reanalyzed_count = 0
        errors = []
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = {executor.submit(analyze_one, c): c for c in candidates_to_reanalyze}
            for future in as_completed(futures):
                try:
                    token, success, msg = future.result(timeout=300)
                    if success:
                        reanalyzed_count += 1
                    else:
                        errors.append(f"Token {token}: {msg}")
                except Exception as e:
                    errors.append(f"Timeout ou erreur: {str(e)}")
        elapsed = time.time() - start_time
        gc.collect()
        return jsonify({
            'message': f'Reanalyse terminee en {elapsed:.1f}s',
            'reanalyzed_count': reanalyzed_count,
            'total_candidates': len(candidates_to_reanalyze),
            'skipped_closed_posts': candidates_skipped,
            'workers_used': MAX_WORKERS,
            'elapsed_seconds': round(elapsed, 1),
            'errors': errors[:10]
        }), 202
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
@app.route('/api/recruteur/reanalyze-poste/<poste>', methods=['POST'])
@jwt_required()
def reanalyze_by_poste(poste):
    if poste not in POSTES:
        return jsonify({'error': f'Poste inconnu: {poste}'}), 400
    if not is_poste_actif(poste):
        return jsonify({
            'error': f'Le poste "{poste}" est cloture. Reanalyse desactivee.',
            'poste': poste,
            'statut': 'cloture',
            'postes_actifs': POSTES_ACTIFS
        }), 403
    try:
        if not supabase:
            return jsonify({'error': 'Supabase non configure'}), 500
        response = supabase.table('candidats').select('*').eq('poste', poste).execute()
        keys = response.data if response.data else []
        if not keys:
            return jsonify({'message': f'Aucune candidature pour le poste "{poste}"'}), 200
        candidates_with_cv = [k for k in keys if k.get('cv_filename')]
        if not candidates_with_cv:
            return jsonify({'message': f'Aucun CV trouve pour le poste "{poste}"'}), 200
        now_iso = datetime.datetime.now().isoformat()
        for data in candidates_with_cv:
            try:
                supabase.table('candidats').update({
                    "analyse_status": "reanalyzing",
                    "reanalyze_trigger": now_iso,
                    "reanalyze_reason": f"Reanalyse manuelle parallele : {poste}"
                }).eq('token', data.get('token')).execute()
            except Exception:
                pass
        def analyze_one(data):
            try:
                token = data.get('token')
                cv_fn = data.get('cv_filename')
                lm_fn = data.get('lettre_filename')
                att_raw = data.get('attestation_filenames', '[]')
                if not cv_fn:
                    return (token, False, "CV manquant")
                run_analysis_for_candidat(token, cv_fn, lm_fn, att_raw, poste, True)
                return (token, True, "OK")
            except Exception as e:
                return (data.get('token'), False, str(e))
        MAX_WORKERS = min(8, len(candidates_with_cv))
        start_time = time.time()
        reanalyzed_count = 0
        errors = []
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = [executor.submit(analyze_one, c) for c in candidates_with_cv]
            for future in as_completed(futures):
                try:
                    token, success, msg = future.result(timeout=300)
                    if success:
                        reanalyzed_count += 1
                    else:
                        errors.append(f"Token {token}: {msg}")
                except Exception as e:
                    errors.append(f"Erreur: {str(e)}")
        elapsed = time.time() - start_time
        gc.collect()
        return jsonify({
            'message': f'Reanalyse terminee pour le poste "{poste}"',
            'poste': poste,
            'statut': 'actif',
            'reanalyzed_count': reanalyzed_count,
            'total_candidates': len(candidates_with_cv),
            'workers_used': MAX_WORKERS,
            'elapsed_seconds': round(elapsed, 1),
            'errors': errors[:10]
        }), 202
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
@app.route('/api/recruteur/cleanup-closed', methods=['POST'])
@jwt_required()
def cleanup_closed_statuses():
    if not supabase:
        return jsonify({'error': 'Supabase non configure'}), 500
    response = supabase.table('candidats').select('token, poste, analyse_status').execute()
    fixed = 0
    for row in (response.data or []):
        if row.get('poste') in POSTES_CLOTURES and row.get('analyse_status') in ('reanalyzing', 'pending'):
            supabase.table('candidats').update({"analyse_status": "completed"}).eq('token', row['token']).execute()
            fixed += 1
    return jsonify({
        'message': f'{fixed} dossier(s) de postes clotures stabilises (scores conserves)',
        'fixed': fixed,
        'postes_concernes': POSTES_CLOTURES
    }), 200
@app.route('/api/recruteur/export/<fmt>', methods=['GET'])
@jwt_required()
def export_candidates(fmt):
    try:
        poste_filter = request.args.get('poste', '')
        statut_filter = request.args.get('statut', '')
        date_start = request.args.get('date_start', '')
        date_end = request.args.get('date_end', '')
        min_score = request.args.get('min_score', type=int)
        if not supabase:
            return jsonify({'error': 'Supabase non configure'}), 500
        response = supabase.table('candidats').select('*').execute()
        all_candidats = response.data if response.data else []
        result = []
        for c in all_candidats:
            c['id'] = c.get('token', '')
            if poste_filter and c.get('poste') != poste_filter:
                continue
            if statut_filter:
                statut = c.get('statut', 'en_attente')
                if statut != statut_filter:
                    continue
            if date_start and c.get('date_candidature'):
                cand_date = c.get('date_candidature', '').split('T')[0]
                if cand_date < date_start:
                    continue
            if date_end and c.get('date_candidature'):
                cand_date = c.get('date_candidature', '').split('T')[0]
                if cand_date > date_end:
                    continue
            if min_score is not None and int(c.get('score', 0)) < min_score:
                continue
            for field in ['score_breakdown', 'flags_eliminatoires', 'analyse_details']:
                if c.get(field):
                    try:
                        c[f'{field}_parsed'] = json.loads(c[field])
                    except Exception:
                        pass
            result.append(c)
        result = sort_candidats(result)
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        poste_suffix = f"_{poste_filter.replace(' ', '_')}" if poste_filter else "_global"
        filename_base = f"rapport{poste_suffix}_{ts}"
        if fmt.lower() == 'excel' or fmt.lower() == 'xlsx':
            buf = generate_excel_report_enhanced(result, poste_filter=poste_filter)
            if not buf:
                return jsonify({'error': 'Erreur generation Excel'}), 500
            return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', as_attachment=True, download_name=f'{filename_base}.xlsx')
        elif fmt.lower() == 'pdf':
            buf = generate_pdf_report_enhanced(result, poste_filter=poste_filter)
            if not buf:
                return jsonify({'error': 'Erreur generation PDF'}), 500
            return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name=f'{filename_base}.pdf')
        elif fmt.lower() == 'csv':
            csv_data = generate_csv_report_enhanced(result, poste_filter=poste_filter)
            return send_file(io.BytesIO(csv_data.encode('utf-8-sig')), mimetype='text/csv', as_attachment=True, download_name=f'{filename_base}.csv')
        elif fmt.lower() in ('word', 'docx'):
            buf = generate_word_report_enhanced(result, poste_filter=poste_filter)
            if not buf:
                return jsonify({'error': 'Erreur generation Word'}), 500
            return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=f'{filename_base}.docx')
        return jsonify({'error': 'Format non supporte. Utilisez: csv, excel, pdf ou word'}), 400
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
@app.route('/api/recruteur/candidats/<token>', methods=['DELETE'])
@jwt_required()
def delete_candidat(token):
    if not supabase:
        return jsonify({'error': 'Supabase non configure'}), 500
    response = supabase.table('candidats').select('*').eq('token', token).execute()
    if not response.data or len(response.data) == 0:
        return jsonify({'error': 'Candidat introuvable'}), 404
    data = response.data[0]
    files_to_delete = []
    if data.get('cv_filename'):
        files_to_delete.append(data.get('cv_filename'))
    if data.get('lettre_filename'):
        files_to_delete.append(data.get('lettre_filename'))
    try:
        attestations = json.loads(data.get('attestation_filenames', '[]'))
        for att in attestations:
            if att:
                files_to_delete.append(att)
    except:
        pass
    deleted_files = []
    failed_files = []
    for filename in files_to_delete:
        try:
            supabase.storage.from_(SUPABASE_STORAGE_BUCKET).remove([filename])
            deleted_files.append(filename)
            logger.info(f"Fichier supprime: {filename}")
        except Exception as e:
            logger.error(f"Erreur suppression fichier {filename}: {e}")
            failed_files.append(filename)
    try:
        supabase.table('candidats').delete().eq('token', token).execute()
        logger.info(f"Candidat {token} supprime avec succes")
    except Exception as e:
        logger.error(f"Erreur suppression candidat {token}: {e}")
        return jsonify({'error': 'Erreur lors de la suppression du candidat', 'details': str(e)}), 500
    return jsonify({
        'message': 'Candidat supprime avec succes',
        'token': token,
        'files_deleted': deleted_files,
        'files_failed': failed_files,
        'candidat': {
            'nom': data.get('nom'),
            'prenom': data.get('prenom'),
            'email': data.get('email'),
            'poste': data.get('poste')
        }
    }), 200
@app.route('/api/recruteur/dossiers/zip/start', methods=['GET'])
@jwt_required()
def start_zip_export():
    _cleanup_old_zip_jobs()
    poste_filter = request.args.get('poste', '')
    date_start = request.args.get('date_start', '')
    date_end = request.args.get('date_end', '')
    job_id = uuid.uuid4().hex
    with _ZIP_JOBS_LOCK:
        _ZIP_JOBS[job_id] = {
            'status': 'pending',
            'created_at': time.time(),
            'progress': 0,
            'total': 0,
            'filepath': None,
            'filename': None,
            'error': None
        }
    threading.Thread(target=_run_zip_export_job, args=(job_id, poste_filter, date_start, date_end), daemon=True).start()
    return jsonify({'job_id': job_id}), 202
@app.route('/api/recruteur/dossiers/zip/status/<job_id>', methods=['GET'])
@jwt_required()
def zip_export_status(job_id):
    with _ZIP_JOBS_LOCK:
        job = _ZIP_JOBS.get(job_id)
        if not job:
            return jsonify({'error': 'Job introuvable ou expire'}), 404
        return jsonify({
            'status': job['status'],
            'progress': job.get('progress', 0),
            'total': job.get('total', 0),
            'error': job.get('error')
        }), 200
@app.route('/api/recruteur/dossiers/zip/download/<job_id>', methods=['GET'])
@jwt_required()
def zip_export_download(job_id):
    with _ZIP_JOBS_LOCK:
        job = _ZIP_JOBS.get(job_id)
        if not job:
            return jsonify({'error': 'Job introuvable ou expire'}), 404
        if job['status'] == 'error':
            return jsonify({'error': job.get('error', 'Erreur inconnue')}), 500
        if job['status'] != 'done':
            return jsonify({'error': 'Export pas encore termine', 'status': job['status']}), 425
        filepath = job.get('filepath')
        if not filepath or not os.path.exists(filepath):
            return jsonify({'error': 'Fichier expire ou deja supprime'}), 410
        response_obj = send_file(filepath, mimetype='application/zip', as_attachment=True, download_name=job['filename'])
        @response_obj.call_on_close
        def _cleanup_after_send():
            try:
                if os.path.exists(filepath):
                    os.remove(filepath)
            except Exception as e:
                logger.warning(f"Nettoyage fichier temporaire ZIP (job {job_id}) echoue: {e}")
            with _ZIP_JOBS_LOCK:
                _ZIP_JOBS.pop(job_id, None)
        return response_obj
def _run_zip_export_job(job_id, poste_filter, date_start, date_end):
    tmp_zip_path = None
    start_time = time.time()
    try:
        with _ZIP_JOBS_LOCK:
            _ZIP_JOBS[job_id]['status'] = 'processing'
            logger.info(f"[job {job_id}] Debut export ZIP optimise - {datetime.datetime.now()}")
        if not supabase:
            with _ZIP_JOBS_LOCK:
                _ZIP_JOBS[job_id]['status'] = 'error'
                _ZIP_JOBS[job_id]['error'] = 'Supabase non configure'
            return
        response = supabase.table('candidats').select('*').execute()
        all_candidats = response.data if response.data else []
        candidats = []
        for c in all_candidats:
            c['id'] = c.get('token', '')
            if poste_filter and c.get('poste') != poste_filter:
                continue
            date_cand = c.get('date_candidature', '')
            if date_cand:
                date_only = date_cand.split('T')[0] if 'T' in date_cand else date_cand[:10]
                if date_start and date_only < date_start:
                    continue
                if date_end and date_only > date_end:
                    continue
            candidats.append(c)
        if not candidats:
            with _ZIP_JOBS_LOCK:
                _ZIP_JOBS[job_id]['status'] = 'error'
                _ZIP_JOBS[job_id]['error'] = 'Aucun dossier a exporter'
            return
        download_tasks = []
        candidats_meta = {}
        for cand in candidats:
            poste_nom = cand.get('poste', 'Poste_Inconnu')
            poste_nom_clean = re.sub(r'[<>:"/\\|?*]', '_', poste_nom)
            num_dossier = cand.get('numero_dossier', '') or f"candidat_{cand['id'][:8]}"
            nom_candidat = cand.get('nom', 'N/A').upper()
            prenom_candidat = cand.get('prenom', 'N/A')
            dossier_candidat_nom = f"{num_dossier} - {nom_candidat} {prenom_candidat}"
            dossier_candidat_nom = re.sub(r'[<>:"/\\|?*]', '_', dossier_candidat_nom)
            dossier_parent = f"{poste_nom_clean}/{dossier_candidat_nom}"
            candidats_meta[cand['id']] = {
                'dossier_parent': dossier_parent,
                'num_dossier': num_dossier,
                'cand': cand,
                'files_written': 0
            }
            cv_file = cand.get('cv_filename', '')
            if cv_file:
                download_tasks.append((cand['id'], cv_file, dossier_parent, 'CV'))
            lettre_file = cand.get('lettre_filename', '')
            if lettre_file:
                download_tasks.append((cand['id'], lettre_file, dossier_parent, 'Lettre_de_motivation'))
            att_raw = cand.get('attestation_filenames', '[]')
            try:
                att_files = json.loads(att_raw) if isinstance(att_raw, str) else att_raw
                for idx, att_file in enumerate(att_files, 1):
                    if att_file:
                        download_tasks.append((cand['id'], att_file, dossier_parent, f'Attestation_{idx}'))
            except Exception:
                pass
        with _ZIP_JOBS_LOCK:
            _ZIP_JOBS[job_id]['total'] = len(download_tasks)
            logger.info(f"[job {job_id}] {len(download_tasks)} fichiers a telecharger")
        def _download_one(task):
            cand_id, blob_name, dossier_parent, prefix = task
            file_bytes = download_file_from_supabase_robust(blob_name)
            return (cand_id, blob_name, dossier_parent, prefix, file_bytes)
        tmp_fd = tempfile.NamedTemporaryFile(delete=False, suffix='.zip', dir='/tmp')
        tmp_zip_path = tmp_fd.name
        tmp_fd.close()
        files_added = 0
        max_workers = min(_ZIP_MAX_WORKERS, max(1, len(download_tasks)))
        with zipfile.ZipFile(tmp_zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            if download_tasks:
                BATCH_SIZE = 50
                total_batches = (len(download_tasks) + BATCH_SIZE - 1) // BATCH_SIZE
                for batch_idx in range(0, len(download_tasks), BATCH_SIZE):
                    batch = download_tasks[batch_idx:batch_idx + BATCH_SIZE]
                    logger.info(f"[job {job_id}] Traitement du lot {batch_idx//BATCH_SIZE + 1}/{total_batches} ({len(batch)} fichiers)")
                    with ThreadPoolExecutor(max_workers=min(max_workers, len(batch))) as executor:
                        futures = [executor.submit(_download_one, t) for t in batch]
                        for future in as_completed(futures):
                            try:
                                cand_id, blob_name, dossier_parent, prefix, file_bytes = future.result(timeout=60)
                            except Exception as e:
                                logger.error(f"[job {job_id}] Erreur telechargement: {e}")
                                continue
                            if file_bytes:
                                ext = blob_name.rsplit('.', 1)[-1].lower() if '.' in blob_name else ''
                                archive_name = f"{dossier_parent}/{prefix}.{ext}" if ext else f"{dossier_parent}/{prefix}"
                                try:
                                    zip_file.writestr(archive_name, file_bytes)
                                    files_added += 1
                                    if cand_id in candidats_meta:
                                        candidats_meta[cand_id]['files_written'] += 1
                                except Exception as e:
                                    logger.error(f"[job {job_id}] Erreur ecriture ZIP: {e}")
                                finally:
                                    del file_bytes
                            with _ZIP_JOBS_LOCK:
                                _ZIP_JOBS[job_id]['progress'] = _ZIP_JOBS[job_id].get('progress', 0) + 1
                            gc.collect()
            for cand_id, meta in candidats_meta.items():
                if meta['files_written'] > 0:
                    continue
                cand = meta['cand']
                info_content = f"""Candidat: {cand.get('nom', 'N/A')} {cand.get('prenom', 'N/A')}
Poste: {cand.get('poste', 'N/A')}
Numero dossier: {meta['num_dossier']}
Email: {cand.get('email', 'N/A')}
Telephone: {cand.get('telephone', 'N/A')}
Date candidature: {cand.get('date_candidature', 'N/A')}"""
                archive_name = f"{meta['dossier_parent']}/INFOS_CANDIDAT.txt"
                zip_file.writestr(archive_name, info_content.encode('utf-8'))
                files_added += 1
        elapsed = time.time() - start_time
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        poste_suffix = f"_{poste_filter.replace(' ', '_')}" if poste_filter else ""
        filename = f"dossiers_candidats{poste_suffix}_{ts}.zip"
        logger.info(f"[job {job_id}] Export ZIP termine en {elapsed:.2f}s pour {len(candidats)} candidats ({files_added} fichiers)")
        with _ZIP_JOBS_LOCK:
            _ZIP_JOBS[job_id]['status'] = 'done'
            _ZIP_JOBS[job_id]['filepath'] = tmp_zip_path
            _ZIP_JOBS[job_id]['filename'] = filename
        del candidats_meta, download_tasks
        gc.collect()
    except Exception as e:
        import traceback
        traceback.print_exc()
        logger.error(f"[job {job_id}] Erreur export ZIP: {e}")
        if tmp_zip_path and os.path.exists(tmp_zip_path):
            try:
                os.remove(tmp_zip_path)
            except Exception:
                pass
        with _ZIP_JOBS_LOCK:
            _ZIP_JOBS[job_id]['status'] = 'error'
            _ZIP_JOBS[job_id]['error'] = str(e)
def _cleanup_old_zip_jobs():
    now = time.time()
    with _ZIP_JOBS_LOCK:
        stale = [jid for jid, j in _ZIP_JOBS.items() if now - j.get('created_at', now) > _ZIP_JOBS_MAX_AGE_SECONDS]
        for jid in stale:
            job = _ZIP_JOBS.pop(jid, None)
            if job and job.get('filepath') and os.path.exists(job['filepath']):
                try:
                    os.remove(job['filepath'])
                except Exception:
                    pass
@app.route('/api/recruteur/candidats/<token>/email-preview', methods=['POST'])
@jwt_required()
def email_preview(token):
    if not supabase:
        return jsonify({'error': 'Supabase non configure'}), 500
    response = supabase.table('candidats').select('*').eq('token', token).execute()
    if not response.data or len(response.data) == 0:
        return jsonify({'error': 'Candidat introuvable'}), 404
    data = response.data[0]
    body = request.get_json(silent=True) or {}
    msg_type = body.get('type', data.get('statut', 'en_attente'))
    nom_c = f"{data.get('prenom', '')} {data.get('nom', '')}".strip()
    poste = data.get('poste', '')
    to_email = data.get('email', '')
    sign = "\nCordialement,\nL'equipe Ressources Humaines\nRecrutBank"
    if msg_type == 'retenu':
        sujet = f"Felicitations – Candidature retenue – {poste}"
        corps = f"Madame, Monsieur {nom_c},\nNous avons le plaisir de vous informer que votre candidature pour le poste de {poste} a ete retenue.\nNous vous contacterons tres prochainement." + sign
    elif msg_type == 'entretien':
        sujet = f"Invitation a un entretien – {poste}"
        corps = f"Madame, Monsieur {nom_c},\nSuite a l'examen de votre candidature pour le poste de {poste}, nous avons le plaisir de vous inviter a un entretien.\nNous prendrons contact avec vous pour convenir d'une date." + sign
    else:
        sujet = f"Reponse a votre candidature – {poste}"
        corps = f"Madame, Monsieur {nom_c},\nNous vous remercions de l'interet que vous portez a notre institution.\nApres examen attentif de votre dossier pour le poste de {poste}, nous avons le regret de vous informer que votre candidature n'a pas ete retenue.\nNous vous encourageons a postuler a nouveau." + sign
    return jsonify({'to': to_email, 'nom': nom_c, 'sujet': sujet, 'corps': corps}), 200
@app.route('/api/recruteur/uploads/<path:filename>', methods=['GET'])
def serve_upload(filename):
    safe = secure_filename(filename.replace('/', '_'))
    if not safe:
        return jsonify({'error': 'Nom de fichier invalide'}), 400
    url = get_signed_url(safe, expiration_minutes=30)
    if not url:
        return jsonify({'error': 'Fichier introuvable'}), 404
    return redirect(url)
@app.route('/api/recruteur/debug/analyse-ia', methods=['POST'])
@jwt_required()
def debug_analyse_ia():
    data = request.get_json(silent=True) or {}
    cv_text = data.get('cv_text', '')
    lettre_text = data.get('lettre_text', '')
    poste = data.get('poste', '')
    if not cv_text or poste not in GRILLE:
        return jsonify({'error': 'cv_text requis et poste doit exister dans GRILLE'}), 400
    result = analyze_cv_with_ia_only(cv_text, lettre_text, [], poste)
    if not result:
        return jsonify({'error': "L'analyse IA a echoue"}), 500
    return jsonify(result), 200
@app.route('/api/test-email', methods=['GET'])
def test_email():
    try:
        to = request.args.get('to', '')
        if not to:
            return jsonify({'error': 'Parametre ?to= requis'}), 400
        ok = send_email(to, 'Test RecrutBank', f'Ceci est un email de test depuis RecrutBank avec {len(ACTIVE_MODELS)} modele(s) IA.')
        return jsonify({'sent': ok}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500
@app.route('/api/health-version', methods=['GET'])
def health_version():
    active_models = [{"name": m["name"], "model": m["model"], "supports_reasoning": m.get("supports_reasoning", False)} for m in ACTIVE_MODELS]
    return jsonify({
        "version": "v13.1-ia-multi-fallback",
        "postes_actifs": POSTES_ACTIFS,
        "postes_count": len(POSTES),
        "analysis_method": "100% IA avec fallback multi-modeles",
        "active_models": active_models,
        "active_models_count": len(ACTIVE_MODELS),
        "max_concurrent_downloads": DOWNLOAD_MAX_CONCURRENT,
        "zip_max_workers": _ZIP_MAX_WORKERS,
        "ia_provider": "OpenRouter",
        "reasoning_enabled": OPENROUTER_REASONING_ENABLED,
        "json_robust_parsing": True,
        "sous_scores_complets": True,
        "eliminatoire_rejet": True,
        "score_conserve": True,
        "score_somme_sous_scores": True,
        "business_rules_stable": True,
        "grilles_disponibles": list(GRILLE.keys()),
        "scoring_max_chef_division": 14,
        "scoring_max_data_analyst": 14,
        "seuils": {"prioritaire": 11, "potentiel_min": 7, "rejet_max": 6},
        "deployed_at": datetime.datetime.now().isoformat(),
        "version_info": "v13.1-ia-multi-fallback - Multi-modeles IA avec fallback automatique"
    }), 200
if __name__ == '__main__':
    port = int(os.getenv("PORT", 10000))
    import multiprocessing
    cpu_count = multiprocessing.cpu_count()
    suggested_workers = min(4, cpu_count * 2)
    logger.info("=" * 60)
    logger.info("🚀 RecrutBank API v13.1 - Multi-modeles IA avec fallback")
    logger.info("=" * 60)
    logger.info(f"Port: {port}")
    logger.info(f"Workers suggeres: {suggested_workers}")
    if ACTIVE_MODELS:
        logger.info(f"✅ {len(ACTIVE_MODELS)} modele(s) IA actif(s):")
        for m in ACTIVE_MODELS:
            supports = "🧠" if m.get("supports_reasoning", False) and OPENROUTER_REASONING_ENABLED else "📝"
            logger.info(f"   {supports} {m['name']} ({m['model']})")
        logger.info(f"Reasoning: {'✅ Active' if OPENROUTER_REASONING_ENABLED else '❌ Desactive'}")
        logger.info(f"Analyse: 100% IA avec fallback automatique")
        logger.info(f"POSTES ACTIFS: {POSTES_ACTIFS}")
        logger.info(f"POSTES CLOTURES: {POSTES_CLOTURES}")
        logger.info(f"GRILLES DISPONIBLES: {len(GRILLE)} postes")
        logger.info(f"Concurrence IA max: {os.getenv('IA_MAX_CONCURRENCY', '5')}")
    else:
        logger.warning("⚠️ AUCUN MODELE IA DISPONIBLE - Le systeme ne peut pas fonctionner")
        logger.warning("   Veuillez configurer OPENROUTER_API_KEY")
    logger.info(f"Telechargements concurrents: {DOWNLOAD_MAX_CONCURRENT}")
    logger.info(f"Workers ZIP max: {_ZIP_MAX_WORKERS}")
    logger.info("=" * 60)
    try:
        import gunicorn
        app.run(host="0.0.0.0", port=port, debug=False, threaded=True)
    except ImportError:
        app.run(host="0.0.0.0", port=port, debug=False, threaded=True)
